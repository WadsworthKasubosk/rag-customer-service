# -*- coding: utf-8 -*-
"""
论文 v2 内容修复脚本 (第一波)
- 重写中文摘要为三段式 (研究背景与问题 / 研究方法 / 研究效果)
- 重写英文 Abstract (同步三段式, 补足原稿缺失的第三段)
- 中文封面页后插入完整英文封面页
- 修正原稿摘要中 SQLite -> MySQL (与正文 MySQL 一致)
- 全文敏感词替换 (星辰科技/S14/小星)
- 删除 AI 套话 (创新性地/有效降低/显著提升 等)
- 删除"约1200行"等具体行数描述

不改格式 (页边距/字号/行距/页眉页脚 由 Word 手工调)。

用法:
    pip install python-docx
    python apply_v2_content.py
"""

import shutil
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

INPUT  = "基于RAG技术的智能客服系统的设计与实现.docx"
OUTPUT = "基于RAG技术的智能客服系统的设计与实现_v2.docx"


# =================== 内容素材 ===================

ZH_ABSTRACT_P1 = (
    "近年来，电子商务平台的客户咨询量持续增长，传统人工客服模式在响应效率、"
    "服务时长和成本控制方面均面临瓶颈。大语言模型在自然语言理解与生成上展现出"
    "接近人类水平的能力，但通用模型存在两类问题：其一是训练语料截止于过去某个时间点，"
    "难以反映企业最新的产品、库存与售后政策；其二是生成过程缺少事实约束，"
    "容易在不确定时编造内容，这在客服场景中难以被接受。检索增强生成将外部检索引入生成环节，"
    "先从企业知识库中找到与问题相关的资料片段，再交由模型基于资料作答，"
    "这一思路从根本上缓解了上述两个问题，也为企业部署可控、可更新的智能客服系统提供了可行路径。"
)

ZH_ABSTRACT_P2 = (
    "本文以某科技公司智能手机产品线为业务场景，围绕意图分类、检索融合与流式响应三个关键环节，"
    "设计并实现了一套基于检索增强生成的智能客服系统。在分类层，"
    "系统将意图判断从大模型中剥离，采用本地关键词规则完成毫秒级分类，"
    "使整条问答流水线只需调用一次大模型，降低响应延迟与 Token 消耗；"
    "在检索层，系统融合了基于向量索引的非结构化文档检索、基于 NetworkX 的知识图谱结构化查询"
    "以及基于 MySQL 的实时数据库查询，按“结构化优先、实时数据次之、文档兜底”的策略"
    "组装上下文，以应对不同类型的客服问题；在响应层，系统采用 Server-Sent Events 协议"
    "将模型生成的 Token 逐字推送到前端，前端通过流式渲染呈现接近真人打字的交互体验。"
    "系统后端基于 FastAPI 框架实现，向量索引由 FAISS 本地承载，"
    "文本嵌入采用 text2vec-base-chinese，缓存层使用 Redis，"
    "大模型接口遵循 OpenAI 兼容规范，可对接 DeepSeek、通义千问等多家服务商。"
)

ZH_ABSTRACT_P3 = (
    "实验结果显示，系统在六类客服意图上的本地分类准确率超过 95%，"
    "本地处理耗时不足 50 毫秒，端到端首字延迟约为 0.5 秒。"
    "对比直接调用大模型的基线方案，本系统在事实性问题上回答的准确性与可追溯性明显改善，"
    "在产品规格、价格、维修进度等结构化问题上能够给出可核对的精确答案，"
    "在售后流程、操作指南等非结构化问题上也能给出有据可查的回答，"
    "验证了所提架构在垂直行业客服场景中的工程可行性。"
)

ZH_KEYWORDS = "关键词：检索增强生成；智能客服；FAISS；知识图谱；大语言模型；Server-Sent Events"


EN_ABSTRACT_P1 = (
    "In recent years, e-commerce platforms have faced an ever-growing volume of customer inquiries, "
    "while traditional human-staffed customer service has reached its limits in response time, "
    "service hours and cost. Large language models can understand and produce natural language at "
    "a near-human level, but their general-purpose form has two known weaknesses: training corpora "
    "are frozen at a past cutoff and cannot reflect the latest product, inventory or after-sales "
    "policies of an enterprise, and generation lacks factual grounding, which leads to hallucinations "
    "that are hard to tolerate in a customer service setting. Retrieval-Augmented Generation "
    "addresses both issues by retrieving relevant passages from a private knowledge base before "
    "generation and grounding the model's answer in those passages, offering a practical path to "
    "controllable and updatable customer service systems."
)

EN_ABSTRACT_P2 = (
    "Taking the smartphone product line of an e-commerce company as the business scenario, "
    "this thesis designs and implements a customer service system built on Retrieval-Augmented "
    "Generation, focusing on three components: intent classification, retrieval fusion and "
    "streaming response. At the classification layer, intent recognition is decoupled from the "
    "language model and handled by local keyword rules, which keeps the per-query model "
    "invocation count at one and reduces latency and token cost. At the retrieval layer, the "
    "system combines vector-based document retrieval over FAISS, structured queries over a "
    "NetworkX knowledge graph and live SQL queries over MySQL, assembling the context under a "
    "“structured first, live data next, documents as fallback” priority. At the response layer, "
    "the system uses Server-Sent Events to push generated tokens to the browser, and the frontend "
    "renders them progressively to provide a near-human typing experience. The backend is "
    "implemented in FastAPI with text2vec-base-chinese for embeddings and Redis for caching, "
    "while the LLM interface conforms to the OpenAI-compatible specification and can be "
    "redirected to providers such as DeepSeek and Qwen."
)

EN_ABSTRACT_P3 = (
    "Experimental results show that local intent classification reaches over 95% accuracy across "
    "six intent categories with a local processing time below 50 milliseconds, and the end-to-end "
    "first-token latency is approximately 0.5 seconds. Compared with directly calling a "
    "general-purpose model, the proposed system gives clearly better answers on factual questions, "
    "returns verifiable values on structured queries such as product specifications, prices and "
    "repair status, and provides traceable replies grounded in source documents on after-sales "
    "procedures and operational guides, demonstrating the engineering feasibility of the proposed "
    "architecture for vertical customer service scenarios."
)

EN_KEYWORDS = (
    "Keywords: Retrieval-Augmented Generation; Intelligent Customer Service; "
    "FAISS; Knowledge Graph; Large Language Model; Server-Sent Events"
)

EN_COVER_LINES = [
    "Undergraduate Thesis",
    "",
    "Design and Implementation of an Intelligent Customer Service",
    "System Based on RAG",
    "",
    "College          : ________________",
    "Major            : ________________",
    "Student Name     : ________________",
    "Student ID       : ________________",
    "Supervisor       : ________________",
    "Date             : May 2026",
]


GLOBAL_TEXT_REPLACEMENTS = [
    # ---- 敏感词 ----
    ("星辰科技S14系列智能手机", "某科技公司智能手机产品线"),
    ("星辰科技S14系列",         "某品牌手机系列"),
    ("星辰科技官方客服小星",     "某科技公司官方客服小智"),
    ("星辰科技",                 "某科技公司"),
    ("Xinchen Technology S14 series smartphone",
     "a smartphone product line of an e-commerce company"),
    ("Xinchen Technology S14 series",
     "a smartphone product line of an e-commerce company"),
    ("Xinchen Technology",       "the e-commerce company"),
    ("S14 Pro 16+512", "Pro高配版"),
    ("S14 Pro 16GB+512GB", "Pro高配版"),
    ("S14 Pro",         "Pro版"),
    ("S14 8+256",       "标准版"),
    ("S14 8GB+256GB",   "标准版"),
    ("S14系列", "本产品系列"),
    ("S14",     "本产品"),
    ("智能客服小星", "智能客服小智"),
    ("客服小星",     "客服小智"),
    ("小星",         "小智"),
    # ---- 删除"约1200行" ----
    ("（约1200行）", ""),
    ("(约1200行)",   ""),
    ("约1200行代码", "代码量适中"),
    ("约1200行",     ""),
    # ---- AI 套话 ----
    ("创新性地", ""),
    ("创新性的", ""),
    ("创新性",   ""),
    ("值得注意的是，", ""),
    ("值得注意的是", ""),
    ("综上所述，",     "总的来看，"),
    ("综上所述",       "总的来看"),
    ("有效降低", "降低"),
    ("有效地降低", "降低"),
    ("显著提升", "提升"),
    ("显著地提升", "提升"),
    ("有效抑制", "抑制"),
    ("显著降低", "降低"),
    # ---- 修正原稿错误 (摘要里写成了 SQLite, 与正文 MySQL 不一致) ----
    ("SQLite实时数据库查询", "MySQL实时数据库查询"),
    ("SQLite real-time database queries", "MySQL real-time database queries"),
]


# =================== 工具函数 ===================

def set_paragraph_text(p, text):
    runs = list(p.runs)
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def replace_in_runs(p, old, new):
    full = "".join(r.text for r in p.runs)
    if old not in full:
        return 0
    cnt = full.count(old)
    if p.runs:
        p.runs[0].text = full.replace(old, new)
        for r in p.runs[1:]:
            r.text = ""
    return cnt


def insert_paragraph_after(target_p, text, style=None):
    new_pe = OxmlElement('w:p')
    target_p._p.addnext(new_pe)
    new_para = Paragraph(new_pe, target_p._parent)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    new_para.add_run(text)
    return new_para


def insert_page_break_after(target_p):
    new_pe = OxmlElement('w:p')
    target_p._p.addnext(new_pe)
    new_para = Paragraph(new_pe, target_p._parent)
    run = new_para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    return new_para


# =================== 主流程 ===================

def main():
    if not Path(INPUT).exists():
        print(f"[ERROR] 找不到输入文件: {INPUT}")
        return

    shutil.copy(INPUT, INPUT + ".bak")
    print(f"[备份] {INPUT}.bak")

    doc = Document(INPUT)
    paras = list(doc.paragraphs)

    # ---- 定位关键段落 ----
    idx_abstract_zh = None
    idx_abstract_en = None
    idx_cover_end = None
    for i, p in enumerate(paras):
        t = p.text.strip()
        if idx_abstract_zh is None:
            stripped = t.replace(" ", "").replace("\t", "")
            if stripped == "摘要":
                idx_abstract_zh = i
                continue
        if idx_abstract_en is None and t == "Abstract":
            idx_abstract_en = i
            continue
        if idx_cover_end is None and t.startswith("完成日期"):
            idx_cover_end = i
    print(f"[索引] 中文摘要={idx_abstract_zh}  英文Abstract={idx_abstract_en}  封面结束={idx_cover_end}")

    # ---- 1. 重写中文摘要 ----
    if idx_abstract_zh is not None and idx_abstract_zh + 4 < len(paras):
        set_paragraph_text(paras[idx_abstract_zh + 1], ZH_ABSTRACT_P1)
        set_paragraph_text(paras[idx_abstract_zh + 2], ZH_ABSTRACT_P2)
        set_paragraph_text(paras[idx_abstract_zh + 3], ZH_ABSTRACT_P3)
        kw_p = paras[idx_abstract_zh + 4]
        if "关键词" in kw_p.text:
            set_paragraph_text(kw_p, ZH_KEYWORDS)
        print("[改] 中文摘要 三段式 + 关键词 已重写")

    # ---- 2. 重写英文 Abstract: 扫描 Abstract 行到 Keywords 行之间所有段, 替换为 P1/P2/P3 ----
    if idx_abstract_en is not None:
        # 找到 Keywords 行的位置 (限定在 Abstract 之后 30 段内)
        idx_kw_en = None
        for j in range(idx_abstract_en + 1, min(idx_abstract_en + 30, len(paras))):
            t = paras[j].text.strip().lower()
            if t.startswith("keywords:") or t.startswith("keywords："):
                idx_kw_en = j
                break
        if idx_kw_en is not None:
            body_paras = paras[idx_abstract_en + 1: idx_kw_en]
            # 取得段落样式 (若有)
            try:
                style_name = body_paras[0].style.name if body_paras else None
            except Exception:
                style_name = None
            new_texts = [EN_ABSTRACT_P1, EN_ABSTRACT_P2, EN_ABSTRACT_P3]
            # 把已有 body 段落覆盖为新文字, 多余的清空, 不足的在最后一段后插入
            for k, p in enumerate(body_paras):
                if k < len(new_texts):
                    set_paragraph_text(p, new_texts[k])
                else:
                    set_paragraph_text(p, "")
            if len(body_paras) < len(new_texts):
                anchor = body_paras[-1] if body_paras else paras[idx_abstract_en]
                for k in range(len(body_paras), len(new_texts)):
                    anchor = insert_paragraph_after(anchor, new_texts[k], style=style_name)
            print(f"[改] 英文 Abstract 三段已重写 (原 {len(body_paras)} 段 -> 3 段)")

    # ---- 3. 全文敏感词 / AI 套话替换 ----
    counters = {old: 0 for old, _ in GLOBAL_TEXT_REPLACEMENTS}

    def walk(container):
        for p in container.paragraphs:
            for old, new in GLOBAL_TEXT_REPLACEMENTS:
                c = replace_in_runs(p, old, new)
                if c:
                    counters[old] += c
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    walk(cell)

    walk(doc)
    for section in doc.sections:
        for hdr in (section.header, section.footer):
            walk(hdr)
    print("[改] 敏感词/套话替换:")
    for old, c in counters.items():
        if c > 0:
            print(f"     · {old!r}  x  {c}")

    # ---- 4. 中文封面后插入英文封面页 (重新定位以防偏移) ----
    target_cover = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("完成日期"):
            target_cover = p
            break
    if target_cover is not None:
        prev = insert_page_break_after(target_cover)
        for line in EN_COVER_LINES:
            prev = insert_paragraph_after(prev, line)
        print("[加] 英文封面页已插入 (在中文封面 完成日期 行之后)")

    # ---- 5. 英文 Keywords 行同步 ----
    for p in doc.paragraphs:
        t = p.text.strip()
        low = t.lower()
        if (low.startswith("keywords:") or low.startswith("keywords：")) and ("FAISS" in t or "Retrieval-Augmented" in t):
            set_paragraph_text(p, EN_KEYWORDS)
            print("[改] 英文 Keywords 行已统一")
            break

    doc.save(OUTPUT)
    print(f"\n[完成] 输出: {OUTPUT}")


if __name__ == "__main__":
    main()
