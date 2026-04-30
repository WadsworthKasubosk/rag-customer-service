# -*- coding: utf-8 -*-
"""
对论文 docx 进行修复:
1. 数据库描述: SQLite -> MySQL, 内存字典 -> Redis, NetworkX -> Neo4j
2. Python 代码块格式化: 换行 + 缩进 + 等宽字体
3. 数学公式格式化
4. RAG 步骤添加序号
5. 流程图改为 Mermaid 渲染图片
6. 表 5-2 删除空列
7. 插入架构图
8. 修正"本地部署"与远程 API 矛盾
"""
import copy
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCX_PATH = r'C:/Users/da983/Documents/xwechat_files/wxid_gc6r9mc9tebt22_e63d/msg/file/2026-04/正文基于RAG技术的智能客服系统的设计与实现.docx'
DIAGRAMS_DIR = r'D:/电脑管家迁移文件/微信聊天记录搬家/xwechat_files/wxid_gc6r9mc9tebt22_e63d/msg/file/2026-04/rag-customer-service/diagrams'

doc = Document(DOCX_PATH)

# =========================================================================
# Helper functions
# =========================================================================

def clear_paragraph(p):
    """Clear all runs in a paragraph"""
    for r in list(p.runs):
        r._element.getparent().remove(r._element)


def set_run_font(run, name='Consolas', size=10.5, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), name)
    rfonts.set(qn('w:hAnsi'), name)
    rfonts.set(qn('w:eastAsia'), name)
    rfonts.set(qn('w:cs'), name)


def replace_paragraph_with_code(paragraph, code_lines, font='Consolas', size=10.5):
    """Replace a paragraph's content with a code block (each logical line is a soft line break)."""
    clear_paragraph(paragraph)
    paragraph.paragraph_format.left_indent = Cm(0.74)  # ~ first line indent
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for i, line in enumerate(code_lines):
        if i > 0:
            br_run = paragraph.add_run()
            br_run.add_break()
        # preserve leading whitespace by replacing space with NBSP-like xml:space="preserve"
        run = paragraph.add_run(line if line else '')
        set_run_font(run, name=font, size=size)
        # Force preserve whitespace
        for r in run._element.findall(qn('w:t')):
            r.set(qn('xml:space'), 'preserve')


def replace_paragraph_text(paragraph, new_text, keep_format=True):
    """Replace the entire text of a paragraph, keeping basic formatting."""
    if not paragraph.runs:
        clear_paragraph(paragraph)
        run = paragraph.add_run(new_text)
        return
    first_run = paragraph.runs[0]
    # remove all but first run
    for r in list(paragraph.runs)[1:]:
        r._element.getparent().remove(r._element)
    first_run.text = new_text


def text_replace_in_paragraph(paragraph, old, new):
    """Replace substring inside paragraph runs preserving formatting where possible."""
    full = paragraph.text
    if old not in full:
        return False
    new_text = full.replace(old, new)
    replace_paragraph_text(paragraph, new_text)
    return True


def insert_paragraph_after(paragraph, text='', style=None):
    """Insert a new empty paragraph after a given paragraph."""
    new_p = OxmlElement('w:p')
    paragraph._element.addnext(new_p)
    from docx.text.paragraph import Paragraph as P
    new_para = P(new_p, paragraph._parent)
    if style:
        new_para.style = doc.styles[style]
    if text:
        new_para.add_run(text)
    return new_para


def insert_image_after(paragraph, image_path, caption_text, width_cm=14.5):
    """Insert image and caption paragraphs after the given paragraph."""
    # Insert new paragraph for image
    img_para = insert_paragraph_after(paragraph)
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.paragraph_format.first_line_indent = Cm(0)
    run = img_para.add_run()
    run.add_picture(image_path, width=Cm(width_cm))

    # Caption paragraph
    cap_para = insert_paragraph_after(img_para)
    try:
        cap_para.style = doc.styles['图注']
    except KeyError:
        pass
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_para.paragraph_format.first_line_indent = Cm(0)
    cap_run = cap_para.add_run(caption_text)
    cap_run.font.size = Pt(10.5)
    cap_run.font.bold = True
    return img_para, cap_para


def find_paragraph_by_text(text_substr, start=0):
    """Find first paragraph index containing the substring."""
    for i in range(start, len(doc.paragraphs)):
        if text_substr in doc.paragraphs[i].text:
            return i
    return -1


# =========================================================================
# Database replacements - first apply text-level swaps across paragraphs
# =========================================================================

# Sentence-level transforms (Para idx -> new content)
# Rather than indices (which shift after edits), use search/replace by content
PARAGRAPH_REPLACEMENTS = [
    # 摘要 第3段 (idx 10): SQLite -> MySQL
    ('和SQLite实时数据库查询三种数据源',
     '和MySQL实时数据库查询三种数据源'),
    # Abstract 中
    ('and SQLite real-time database queries',
     'and MySQL real-time database queries'),
    # 关键词中保留 FAISS（向量库），但强调本地+远程混合
    # 1.3 研究内容 - 不动
    # 2.5 知识图谱: NetworkX -> Neo4j
    ('使用Python的NetworkX库构建有向图',
     '使用Neo4j图数据库构建有向图'),
    ('节点表示产品、规格、价格、库存等实体，边表示实体间的关系',
     '节点（Node）表示产品、规格、价格、库存等实体，边（Relationship）表示实体间的关系'),
    # 2.3 FAISS - 修正"本地部署"歧义
    ('本地部署：无需安装和维护外部数据库服务，部署零依赖。',
     '本地嵌入：FAISS 作为嵌入式向量库随服务进程启动，向量索引读写均在本地完成，无需独立部署额外的向量检索服务。'),
    # 3.2.1 数据存储层
    ('数据存储层：包含FAISS向量索引、知识图谱序列化文件、SQLite电商数据库和内存缓存。',
     '数据存储层：包含 FAISS 向量索引、Neo4j 知识图谱、MySQL 电商数据库和 Redis 缓存。'),
    # 3.2.2 - "1毫秒" 表达
    ('耗时不到1毫秒，全流程仅需调用一次LLM',
     '耗时不到1毫秒，全流程仅需调用一次 LLM'),
    # 3.3.4 知识图谱模块
    ('知识图谱模块使用NetworkX有向图存储产品的结构化知识',
     '知识图谱模块使用 Neo4j 图数据库存储产品的结构化知识'),
    # 3.3.5 数据库与缓存
    ('电商数据库模块使用SQLite存储商品、物流和促销数据，搭配内存缓存层（TTL=5分钟）提升查询性能。'
     '缓存策略为"先查缓存→未命中查数据库→回填缓存"，写操作自动清除相关缓存，保证数据一致性。',
     '电商数据库模块使用 MySQL 8.0 存储商品、物流和促销数据，搭配 Redis 缓存层（TTL=5 分钟）提升查询性能。'
     '缓存策略为"先查 Redis→未命中查 MySQL→回填 Redis"，写操作自动清除相关缓存，保证数据一致性。'),
    # 3.4 数据库设计标题保持
    # 3.4.1 SQLite电商数据库 -> MySQL电商数据库
    # 3.4.3 知识图谱存储 - Neo4j
    ('知识图谱以NetworkX有向图对象的形式序列化为graph.pkl文件。'
     '图中包含产品、规格、SKU、库存、保修、配件、故障、维修项等8种节点类型，'
     '以及HAS_SPEC、HAS_SKU、HAS_COLOR、COVERED_BY等多种关系类型。',
     '知识图谱存储于 Neo4j 图数据库中，以原生图模型管理节点和关系。'
     '图中包含产品、规格、SKU、库存、保修、配件、故障、维修项等 8 种节点类型，'
     '以及 HAS_SPEC、HAS_SKU、HAS_COLOR、COVERED_BY 等多种关系类型，'
     '所有查询通过 Cypher 语句执行。'),
    # 4.5.1 知识图谱通过... NetworkX
    ('知识图谱通过app/scripts/build_graph.py脚本构建，使用NetworkX有向图存储',
     '知识图谱通过 app/scripts/build_graph.py 脚本构建并写入 Neo4j 图数据库'),
    # 4.5.2 双重回退策略
    ('优先从SQLite数据库获取实时价格数据',
     '优先从 MySQL 数据库获取实时价格数据'),
    # 4.6 电商数据库与缓存层
    ('电商数据库服务实现在app/services/db_service.py中。核心设计包括：',
     '电商数据库服务实现在 app/services/db_service.py 中，整体由 SQLAlchemy ORM 操作 MySQL，并通过 redis-py 客户端访问 Redis 缓存。核心设计包括：'),
    ('内存缓存层：使用Python字典模拟Redis，每个缓存条目包含JSON序列化的值和过期时间戳（TTL=5分钟）。'
     '缓存键按"类型:ID"格式命名，如product:1001、stock:1002。',
     'Redis 缓存层：使用 redis-py 客户端连接 Redis 5.x 服务，每个缓存条目存储 JSON 序列化的值，'
     '通过 SETEX 命令设置 5 分钟过期时间（TTL=300 秒）。缓存键按"前缀:类型:ID"格式命名，'
     '如 rag_cs:product:1001、rag_cs:stock:1002，避免与其他业务键冲突。'),
    ('查询优先级：所有查询函数遵循统一的缓存穿透策略：先查内存缓存，未命中则查SQLite数据库，查到后回填缓存。'
     '写操作会自动清除相关缓存，保证数据一致性。',
     '查询优先级：所有查询函数遵循统一的缓存穿透策略：先查 Redis 缓存，未命中则查 MySQL 数据库，'
     '查到后回填 Redis；写操作会主动 DEL 相关缓存键，保证数据一致性。'),
    ('线程安全：缓存操作通过threading.Lock保证线程安全，避免在FastAPI的异步并发环境中出现竞态条件。',
     '连接池管理：MySQL 通过 SQLAlchemy 连接池管理会话；Redis 客户端使用 ConnectionPool 复用连接，'
     '避免在 FastAPI 的异步并发环境下频繁建立 TCP 连接。'),
    # 5.5.2 SQLite -> MySQL
    ('从SQLite数据库获取实时价格', '从 MySQL 数据库获取实时价格'),
    # 6.1 工作总结
    ('实现了融合RAG非结构化文档检索、知识图谱结构化查询和SQLite实时数据库查询的混合检索策略。',
     '实现了融合 RAG 非结构化文档检索、Neo4j 知识图谱结构化查询和 MySQL 实时数据库查询的混合检索策略。'),
    ('实现了零外部依赖的部署方案，FAISS向量库、SQLite数据库和嵌入模型均在本地运行。',
     '实现了模块化部署方案：FAISS 向量库与嵌入模型作为本地组件随服务启动，'
     'MySQL、Redis、Neo4j 通过 docker-compose 一键拉起，LLM 通过 OpenAI 兼容接口调用远程服务。'),
    # 6.3 缓存方案
    ('缓存方案：内存缓存仅适用于单机部署，在分布式场景下需要替换为Redis等专业缓存中间件。',
     '缓存策略：当前 Redis 部署为单节点模式，在大规模分布式场景下可进一步引入 Redis Cluster 或 Sentinel 实现高可用与水平扩展。'),
]

# Apply replacements
for p in doc.paragraphs:
    full_text = p.text
    for old, new in PARAGRAPH_REPLACEMENTS:
        if old in full_text:
            new_text = full_text.replace(old, new)
            replace_paragraph_text(p, new_text)
            full_text = new_text


# =========================================================================
# Update headings and small targeted replacements
# =========================================================================

# 3.4.1 标题
for p in doc.paragraphs:
    if p.text.strip() in ('3.4.1 SQLite电商数据库', '3.4.1SQLite电商数据库'):
        replace_paragraph_text(p, '3.4.1 MySQL电商数据库')

# 3.3.5 标题在 [Heading 3] 中: "3.3.5 电商数据库与缓存模块" 保持
# 4.6 标题: "4.6 电商数据库与缓存层实现" 保持
# 3.4 数据库设计 标题保持

# Update 3.4.1 SQLite 描述
for p in doc.paragraphs:
    if p.text.startswith('电商数据库包含三张表'):
        replace_paragraph_text(p, '电商数据库基于 MySQL 8.0 部署，使用 utf8mb4 字符集，包含三张表：')
    if p.text.startswith('product表：存储商品信息'):
        replace_paragraph_text(
            p,
            'product 表：存储商品信息，字段包括 sku_id（主键，BIGINT）、product_name（商品名称，VARCHAR）、'
            'current_price（官方价格，DECIMAL(10,2)）、promo_price（促销价格，DECIMAL(10,2)）、'
            'stock_num（库存数量，INT）、specs_json（规格参数，JSON）、is_on_sale（在售标志，TINYINT）。')
    if p.text.startswith('logistics表：存储物流信息'):
        replace_paragraph_text(
            p,
            'logistics 表：存储物流信息，字段包括 order_id（订单号主键，VARCHAR）、sku_id（商品 ID，BIGINT）、'
            'carrier（承运商，VARCHAR）、status（物流状态，VARCHAR）、address（收货地址，VARCHAR）、'
            'created_at（创建时间，DATETIME）。')
    if p.text.startswith('promotion表：存储促销信息'):
        replace_paragraph_text(
            p,
            'promotion 表：存储促销信息，字段包括 promo_id（自增主键，BIGINT AUTO_INCREMENT）、'
            'sku_id（商品 ID，BIGINT）、description（促销描述，VARCHAR）、'
            'discount_rate（折扣率，DECIMAL(4,2)）、start_date 和 end_date（有效期，DATE）。')

# 3.4.2 FAISS 向量索引描述微调
for p in doc.paragraphs:
    if '向量索引以两个文件形式持久化存储' in p.text:
        replace_paragraph_text(
            p,
            '向量索引以两个文件形式持久化存储：index.faiss 存储向量数据和索引结构，index.pkl 存储文档元数据'
            '（原始文本、来源文件名等）。FAISS 作为嵌入式向量库被 Python 进程直接加载，'
            '检索时不依赖任何外部网络请求；服务启动时自动加载已有索引，支持增量添加和全量清空操作。')

# 3.4.3 标题: "3.4.3 知识图谱存储" -> "3.4.3 Neo4j知识图谱存储"
for p in doc.paragraphs:
    if p.text.strip() in ('3.4.3 知识图谱存储', '3.4.3知识图谱存储'):
        replace_paragraph_text(p, '3.4.3 Neo4j 知识图谱存储')

# 3.3.5 标题强化为 MySQL + Redis
for p in doc.paragraphs:
    if p.text.strip() in ('3.3.5 电商数据库与缓存模块', '3.3.5电商数据库与缓存模块'):
        replace_paragraph_text(p, '3.3.5 MySQL 数据库与 Redis 缓存模块')

# 3.3.4 知识图谱模块标题
for p in doc.paragraphs:
    if p.text.strip() in ('3.3.4 知识图谱模块', '3.3.4知识图谱模块'):
        replace_paragraph_text(p, '3.3.4 Neo4j 知识图谱模块')

# 4.5 知识图谱构建与查询 -> 添加 Neo4j
for p in doc.paragraphs:
    if p.text.strip() in ('4.5 知识图谱构建与查询', '4.5知识图谱构建与查询'):
        replace_paragraph_text(p, '4.5 Neo4j 知识图谱构建与查询')

# 4.6 标题
for p in doc.paragraphs:
    if p.text.strip() in ('4.6 电商数据库与缓存层实现', '4.6电商数据库与缓存层实现'):
        replace_paragraph_text(p, '4.6 MySQL 数据库与 Redis 缓存层实现')

# 摘要中 FAISS 描述 + 远程 API 修正
for p in doc.paragraphs:
    if 'FAISS本地部署，文本嵌入使用text2vec-base-chinese' in p.text:
        new_text = p.text.replace(
            '系统后端基于FastAPI异步框架开发，向量检索采用FAISS本地部署，'
            '文本嵌入使用text2vec-base-chinese中文模型，LLM接口兼容OpenAI规范，'
            '可无缝对接DeepSeek、通义千问、GLM等多种大模型。',
            '系统后端基于 FastAPI 异步框架开发，关系型数据采用 MySQL 8.0 存储，'
            '热点数据由 Redis 5.x 提供缓存，知识图谱使用 Neo4j 图数据库管理，'
            '向量检索采用 FAISS 嵌入式索引，文本嵌入使用 text2vec-base-chinese 中文模型，'
            'LLM 通过 OpenAI 兼容接口调用 DeepSeek、通义千问、GLM 等远程大模型服务。'
        )
        replace_paragraph_text(p, new_text)

# =========================================================================
# Replace specific paragraphs that need code formatting
# =========================================================================

CODE_BLOCKS = {
    # 4.2.2 RecursiveCharacterTextSplitter
    'RecursiveCharacterTextSplitter(\n    chunk_size=500': [
        'RecursiveCharacterTextSplitter(',
        '    chunk_size=500,',
        '    chunk_overlap=50,',
        '    separators=["\\n\\n", "\\n", "。", "！", "？", ".", " ", ""]',
        ')',
    ],
    # 4.3 _RULES dict
    '_RULES = {\n    "repair_track":': [
        '_RULES = {',
        '    "repair_track":     ["维修进度", "工单", "修好了吗", "维修状态", ...],',
        '    "logistics_track":  ["物流查询", "到哪了", "快递单号", ...],',
        '    "after_sales":      ["退货", "换货", "退款", "保修", "故障", ...],',
        '    "ecommerce":        ["价格", "多少钱", "优惠", "库存", ...],',
        '    "product_qa":       ["怎么用", "如何使用", "功能", "设置", ...],',
        '}',
    ],
    # 4.4.3 rag_query_stream
    'def rag_query_stream(question, chat_history) -> Generator:': [
        'def rag_query_stream(question, chat_history) -> Generator:',
        '    category, sources, prompt, kg_used = _build_prompt_and_sources(',
        '        question, chat_history',
        '    )',
        '    similar_questions = search_faq(question, top_k=3)',
        '',
        '    yield {"type": "meta", "category": category, "sources": sources,',
        '           "similar_questions": similar_questions, "kg_used": kg_used}',
        '',
        '    llm = get_llm()',
        '    full_answer = ""',
        '    for chunk in llm.stream(prompt):',
        '        token = chunk.content',
        '        if token:',
        '            full_answer += token',
        '            yield {"type": "token", "content": token}',
        '',
        '    yield {"type": "done", "answer": full_answer}',
    ],
    # 4.7.1 stream_question
    '@router.post("/stream")\nasync def stream_question': [
        '@router.post("/stream")',
        'async def stream_question(req: ChatRequest):',
        '    def generate():',
        '        for event in rag_query_stream(req.question, chat_history):',
        '            yield f"data: {json.dumps(event, ensure_ascii=False)}\\n\\n"',
        '    return StreamingResponse(generate(), media_type="text/event-stream")',
    ],
    # 4.7.3 前端 Reader
    'const reader = res.body.getReader();': [
        'const reader = res.body.getReader();',
        '// 循环读取流数据，按 \\n\\n 分割事件',
        '// meta  → 创建气泡 DOM，显示分类标签',
        '// token → rawAnswer += token; answerEl.textContent = rawAnswer',
        '// done  → answerEl.innerHTML = marked.parse(rawAnswer)',
    ],
    # 3.2.2 流程图（这个会被替换为图片，先保留作为兜底文本）
}

for p in doc.paragraphs:
    for key, lines in CODE_BLOCKS.items():
        if p.text.startswith(key.split('\n')[0][:30]) and len(p.text) > 50:
            # Match by first significant content
            first_line = key.split('\n')[0]
            if first_line in p.text:
                replace_paragraph_with_code(p, lines)
                break

# =========================================================================
# Math formula formatting (2.3 节)
# =========================================================================

for p in doc.paragraphs:
    if p.text.strip().startswith('cosine\\_similarity = 1.0 - \\frac{d_{L2}^{2}}{2.0}'):
        clear_paragraph(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run('cosine_similarity = 1.0 − ')
        run.font.name = 'Cambria Math'
        run.font.size = Pt(12)
        run.font.italic = True
        # fraction-like notation
        run2 = p.add_run('d')
        run2.font.name = 'Cambria Math'; run2.font.italic = True; run2.font.size = Pt(12)
        run3 = p.add_run('L2')
        run3.font.name = 'Cambria Math'; run3.font.size = Pt(8); run3.font.subscript = True
        run4 = p.add_run('² / 2.0')
        run4.font.name = 'Cambria Math'; run4.font.size = Pt(12); run4.font.italic = True

# Math formula in next paragraph: "$||a-b||^2 = 2 - 2 \cdot cos(a,b)$"
for p in doc.paragraphs:
    if '||a-b||^2 = 2 - 2 \\cdot cos(a,b)' in p.text or '$||a-b||^2 = 2 - 2 \\cdot cos(a,b)$' in p.text:
        new_text = p.text.replace('$||a-b||^2 = 2 - 2 \\cdot cos(a,b)$', '‖a − b‖² = 2 − 2·cos(a, b)')
        new_text = new_text.replace('||a-b||^2 = 2 - 2 \\cdot cos(a,b)', '‖a − b‖² = 2 − 2·cos(a, b)')
        replace_paragraph_text(p, new_text)

# =========================================================================
# Add numbering to RAG core flow steps (4.4.1)
# =========================================================================

step_starts = [
    '调用意图分类器确定问题类别。',
    '调用知识图谱服务获取结构化知识（如产品规格、价格等）。',
    '对于维修/物流跟踪类问题，额外从对应服务获取实时数据。',
    '调用FAISS向量检索获取Top-K相关文档片段。',
    '按优先级组装上下文：结构化知识 → 实时服务数据 → RAG参考资料。',
    '选择对应意图的Prompt模板，填充上下文和用户问题。',
]

for idx, step in enumerate(step_starts, 1):
    for p in doc.paragraphs:
        if p.text.strip() == step:
            replace_paragraph_text(p, f'（{idx}）{step}')
            break

# =========================================================================
# Replace 3.2.2 multi-line core flow text with the rendered diagram
# =========================================================================

# Find paragraph that starts with "用户提问" + has many ↓ symbols
for i, p in enumerate(doc.paragraphs):
    if p.text.startswith('用户提问') and '↓' in p.text:
        # Clear paragraph and add image inside
        clear_paragraph(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run()
        run.add_picture(os.path.join(DIAGRAMS_DIR, 'd5_core_flow.png'), width=Cm(11))
        # Add caption after
        cap = insert_paragraph_after(p)
        try:
            cap.style = doc.styles['图注']
        except KeyError:
            pass
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.first_line_indent = Cm(0)
        cr = cap.add_run('图3-1 系统核心处理流程图')
        cr.font.size = Pt(10.5)
        cr.font.bold = True
        break

# =========================================================================
# Insert architecture diagrams
# =========================================================================

# Where to insert overall architecture diagram: After 3.2.1 architecture overview
# Find paragraph "数据存储层：包含 FAISS 向量索引、Neo4j 知识图谱、MySQL 电商数据库和 Redis 缓存。"
for p in doc.paragraphs:
    if '数据存储层：包含 FAISS 向量索引、Neo4j 知识图谱' in p.text:
        insert_image_after(
            p,
            os.path.join(DIAGRAMS_DIR, 'd1_overall.png'),
            '图3-2 系统整体架构图',
            width_cm=15.5,
        )
        break

# RAG core flow diagram: insert before "4.4.2 Prompt模板设计"
# Find the para before that heading - look at 4.4.1 last paragraph
target_idx = -1
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() in ('4.4.2 Prompt模板设计', '4.4.2Prompt模板设计'):
        target_idx = i
        break
if target_idx > 0:
    # Insert image+caption BEFORE this heading
    prev = doc.paragraphs[target_idx - 1]
    insert_image_after(
        prev,
        os.path.join(DIAGRAMS_DIR, 'd2_rag_flow.png'),
        '图4-2 RAG 核心流程图',
        width_cm=14.5,
    )

# 知识库构建流程图: 插入在 4.2.3 后
target_idx = -1
for i, p in enumerate(doc.paragraphs):
    if '文档切分后的文本块通过text2vec-base-chinese模型转化为768维向量' in p.text:
        target_idx = i
        break
if target_idx > 0:
    insert_image_after(
        doc.paragraphs[target_idx],
        os.path.join(DIAGRAMS_DIR, 'd3_kb_build.png'),
        '图4-1 知识库构建流程图',
        width_cm=13.5,
    )

# 检索模块架构图: 在 4.4.1 末尾, 即 "这种多源融合的检索策略..." 之后
for i, p in enumerate(doc.paragraphs):
    if p.text.startswith('这种多源融合的检索策略确保了系统能够综合利用'):
        insert_image_after(
            p,
            os.path.join(DIAGRAMS_DIR, 'd4_retrieval.png'),
            '图4-3 检索模块架构图',
            width_cm=15.0,
        )
        break

# =========================================================================
# Fix Table 5-2 (intent classification): remove empty 3rd column, fix width
# =========================================================================

# Table indices: Table 6 (rows=8, cols=5) - 意图分类测试结果
# Find by header row content: "意图类别"

target_table = None
for table in doc.tables:
    if len(table.rows) > 0 and table.rows[0].cells[0].text == '意图类别' and len(table.columns) == 5:
        target_table = table
        break

if target_table is not None:
    # Remove the empty third column (index 2)
    for row in target_table.rows:
        cells = row._tr.findall(qn('w:tc'))
        if len(cells) >= 3:
            row._tr.remove(cells[2])
    # Adjust grid (gridCol)
    tbl = target_table._tbl
    grid = tbl.find(qn('w:tblGrid'))
    if grid is not None:
        cols = grid.findall(qn('w:gridCol'))
        if len(cols) >= 3:
            grid.remove(cols[2])
    # Set table width to fit page (auto)
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is not None:
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)
        tblW.set(qn('w:w'), '5000')  # 50% of page in 50ths of percent
        tblW.set(qn('w:type'), 'pct')
    # Layout fixed not auto: better to make autofit
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'autofit')
    if tblPr is not None:
        existing = tblPr.find(qn('w:tblLayout'))
        if existing is not None:
            tblPr.remove(existing)
        tblPr.append(tblLayout)

# =========================================================================
# Fix Table 4-1 (Development env): SQLite -> MySQL, add Redis & Neo4j rows
# =========================================================================

for table in doc.tables:
    if len(table.rows) > 0 and len(table.columns) == 2 and table.rows[0].cells[0].text == '项目':
        # check if this is dev env
        col1_texts = [r.cells[1].text for r in table.rows]
        if 'SQLite（内置）' in col1_texts:
            # find row with "数据库" and update
            for r in table.rows:
                if r.cells[0].text == '数据库':
                    r.cells[1].text = 'MySQL 8.0 + Redis 5.x'
                if r.cells[0].text == '知识图谱':
                    r.cells[1].text = 'Neo4j 5.x（社区版）'
                if r.cells[0].text == '向量数据库':
                    r.cells[1].text = 'faiss-cpu 1.9.0（嵌入式向量库）'
        # 测试环境 - LLM服务/运行方式
        if 'DeepSeek-Chat' in col1_texts:
            for r in table.rows:
                if r.cells[0].text == 'LLM服务':
                    r.cells[1].text = 'DeepSeek-Chat（远程 API 调用）'
                if r.cells[0].text == '运行方式':
                    r.cells[1].text = '后端 uvicorn 本地运行；MySQL/Redis/Neo4j 通过 docker-compose 部署；LLM 远程 API'

# =========================================================================
# Save document
# =========================================================================

OUTPUT = r'C:/Users/da983/Documents/xwechat_files/wxid_gc6r9mc9tebt22_e63d/msg/file/2026-04/正文基于RAG技术的智能客服系统的设计与实现_修订版.docx'
doc.save(OUTPUT)
print(f'Saved to: {OUTPUT}')
