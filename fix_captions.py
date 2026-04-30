# -*- coding: utf-8 -*-
"""
修正图片标题编号，保持顺序 = 文档中出现顺序。
"""
from docx import Document

DOCX = r'C:/Users/da983/Documents/xwechat_files/wxid_gc6r9mc9tebt22_e63d/msg/file/2026-04/正文基于RAG技术的智能客服系统的设计与实现_修订版.docx'
doc = Document(DOCX)


def replace_paragraph_text(paragraph, new_text):
    if not paragraph.runs:
        run = paragraph.add_run(new_text)
        return
    first_run = paragraph.runs[0]
    for r in list(paragraph.runs)[1:]:
        r._element.getparent().remove(r._element)
    first_run.text = new_text


# Targeted caption renames (old text -> new text)
RENAMES = {
    '图3-2 系统整体架构图': '图3-1 系统整体架构图',
    '图3-1 系统核心处理流程图': '图3-2 系统核心处理流程图',
    '图4-1 知识库构建流程图': '图4-1 知识库构建流程图',  # keep
    '图4-3 检索模块架构图': '图4-2 检索模块架构图',
    '图4-2 RAG 核心流程图': '图4-3 RAG 核心流程图',
    '图4-1 知识图谱实体关系图': '图4-4 知识图谱实体关系图',
    '图4-4 系统前端界面': '图4-5 系统前端界面',
    '图4-5 知识图谱可视化界面': '图4-6 知识图谱可视化界面',
}

# Apply renames - need to do in two steps to avoid clobbering
TEMP_PREFIX = '__TMP__'
for p in doc.paragraphs:
    text = p.text.strip()
    if text in RENAMES:
        replace_paragraph_text(p, TEMP_PREFIX + RENAMES[text])

for p in doc.paragraphs:
    text = p.text.strip()
    if text.startswith(TEMP_PREFIX):
        replace_paragraph_text(p, text[len(TEMP_PREFIX):])

OUTPUT = DOCX
doc.save(OUTPUT)
print('Captions fixed')
