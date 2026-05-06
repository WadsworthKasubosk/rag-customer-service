# -*- coding: utf-8 -*-
"""
论文内容修复脚本（只改内容，不改格式）
用法:
    pip install python-docx
    python fix_content.py
输入: 基于RAG技术的智能客服系统.docx
输出: 基于RAG技术的智能客服系统_fixed.docx
日志: fix_content.log
"""

import re
import shutil
from pathlib import Path
from docx import Document

INPUT  = "基于RAG技术的智能客服系统的设计与实现.docx"
OUTPUT = "基于RAG技术的智能客服系统的设计与实现_fixed.docx"
LOG    = "fix_content.log"

# ============================================================
# 1. 字符串级替换规则（按顺序执行，长串先替换避免被短串截胡）
# ============================================================
TEXT_REPLACEMENTS = [
    # ---- A1 敏感词 ----
    ("星辰科技S14系列智能手机", "某科技公司智能手机产品线"),
    ("星辰科技S14系列",         "某品牌手机系列"),
    ("星辰科技官方客服小星",     "某科技公司官方客服小智"),
    ("星辰科技",                 "某科技公司"),
    ("Xinchen Technology S14 series smartphone",
     "a smartphone product line of an e-commerce company"),
    ("Xinchen Technology",       "the e-commerce company"),
    # SKU 型号细化（顺序：长在前）
    ("S14 Pro 16+512", "Pro高配版"),
    ("S14 Pro 16GB+512GB", "Pro高配版"),
    ("S14 Pro",         "Pro版"),
    ("S14 8+256",       "标准版"),
    ("S14 8GB+256GB",   "标准版"),
    # 单独 S14 放最后，避免误伤 S14 Pro
    ("S14系列", "本产品系列"),
    ("S14",     "本产品"),
    # 客服昵称
    ("智能客服小星", "智能客服小智"),
    ("客服小星",     "客服小智"),
    ("小星",         "小智"),

    # ---- A3 删除代码行数描述 ----
    ("（约1200行）", ""),
    ("(约1200行)",   ""),
    ("约1200行代码", "代码量适中"),
    ("约1200行",     ""),

    # ---- G1 删除/弱化 AI 套话 ----
    ("创新性地", ""),
    ("值得注意的是，", ""),
    ("值得注意的是", ""),
    ("综上所述，",     "总的来看，"),
    ("综上所述",       "总的来看"),
    ("有效降低", "降低"),
    ("有效地降低", "降低"),
    ("显著提升", "提升"),
    ("显著地提升", "提升"),
    ("有效抑制", "抑制"),
    ("显著降低", "降低"),
]

# ============================================================
# 2. 缩写首次出现补全（只补全文档中第一次出现的位置）
# ============================================================
ABBREV_FIRST_OCCURRENCE = [
    ("KG",   "知识图谱（Knowledge Graph，KG）"),
    ("SKU",  "SKU（Stock Keeping Unit，库存单位）"),
    ("TTFB", "首字响应延迟（Time To First Byte，TTFB）"),
    ("ORM",  "对象关系映射（Object-Relational Mapping，ORM）"),
    ("FAQ",  "常见问题（Frequently Asked Questions，FAQ）"),
]

# ============================================================
# 3. 图编号重排（按段落文本里的"图 X-Y"做映射）
# ============================================================
FIG_RENUMBER = {
    "图3-2 系统总体架构图":             "图3-1 系统总体架构图",
    "图 3-2 系统总体架构图":            "图 3-1 系统总体架构图",
    "图3-2":                            "__TMP_FIG_3_1__",
    "图 3-2":                           "__TMP_FIG_3_1__",
    "图3-1 RAG系统核心处理流程图":       "图3-2 RAG系统核心处理流程图",
    "图 3-1 RAG系统核心处理流程图":      "图 3-2 RAG系统核心处理流程图",
    "图3-1 RAG":                         "图3-2 RAG",
    "图 3-1 RAG":                        "图 3-2 RAG",
    "图4-4 知识图谱实体关系图":          "图3-4 知识图谱实体关系图",
    "图 4-4 知识图谱实体关系图":         "图 3-4 知识图谱实体关系图",
    "图4-4":                             "图3-4",
    "图 4-4":                            "图 3-4",
    "图3-4 Redis缓存穿透策略流程图":     "图3-5 Redis缓存穿透策略流程图",
    "图 3-4 Redis缓存穿透策略流程图":    "图 3-5 Redis缓存穿透策略流程图",
    "图3-5 意图分类器算法流程图":         "图3-6 意图分类器算法流程图",
    "图 3-5 意图分类器算法流程图":        "图 3-6 意图分类器算法流程图",
    "图4-5 系统前端聊天界面":             "图4-4 系统前端聊天界面",
    "图 4-5 系统前端聊天界面":            "图 4-4 系统前端聊天界面",
    "图4-6 知识图谱可视化":                "图4-5 知识图谱可视化",
    "图 4-6 知识图谱可视化":               "图 4-5 知识图谱可视化",
    "__TMP_FIG_3_1__":                    "图 3-1",
}

# ============================================================
# 4. 代码块替换为占位符
# ============================================================
CODE_MARKERS = [
    "def ", "class ", "import ", "from ", "@router", "async def",
    "-> ", "yield ", "return ", "self.", ".append(", ".format(",
    "= None", "if __name__", "raise ", "except ", "try:",
    "const ", "let ", "await ", "function(",
]

CODE_BLOCKS = [
    ("# app/models/llm.py",          "code_2_2_get_llm.png"),
    ("def get_llm()",                 "code_2_2_get_llm.png"),
    ("# app/core/vector_store.py（向量化初始化）", "code_2_4_embedding.png"),
    ("HuggingFaceEmbeddings(",        "code_2_4_embedding.png"),
    ("# app/store/database.py",       "code_2_6_session.png"),
    ("def session_scope()",           "code_2_6_session.png"),
    ("# app/store/cache.py",          "code_2_6_redis.png"),
    ("def get_cache_client()",        "code_2_6_redis.png"),
    ("# app/knowledge/loader.py",     "code_4_2_loader.png"),
    ("def load_document(",            "code_4_2_loader.png"),
    ("# app/core/vector_store.py（知识库构建）", "code_4_2_index.png"),
    ("def build_index(",              "code_4_2_index.png"),
    ("# app/core/classifier.py",      "code_4_3_classify.png"),
    ("def classify_question(",        "code_4_3_classify.png"),
    ("# app/core/rag_chain.py（_build_prompt", "code_4_4_rag.png"),
    ("def _build_prompt_and_sources(", "code_4_4_rag.png"),
    ("# app/core/rag_chain.py（流式生成器）", "code_4_4_stream.png"),
    ("def rag_query_stream(",         "code_4_4_stream.png"),
    ("# app/api/chat.py",             "code_4_5_sse.png"),
    ("async def stream_question(",    "code_4_5_sse.png"),
    ("// 前端 JavaScript",             "code_4_5_frontend.png"),
    ("const res = await fetch",        "code_4_5_frontend.png"),
    ("# app/services/db_service.py",   "code_4_6_db.png"),
    ("def get_product(",               "code_4_6_db.png"),
]


def is_code_paragraph(text: str) -> bool:
    if not text.strip():
        return False
    if text.startswith("    ") or text.lstrip().startswith(("# ", "//", "from ", "import ")):
        if any(m in text for m in CODE_MARKERS):
            return True
    hits = sum(1 for m in CODE_MARKERS if m in text)
    return hits >= 2


def replace_in_runs(paragraph, old: str, new: str) -> int:
    full = "".join(r.text for r in paragraph.runs)
    if old not in full:
        return 0
    count = full.count(old)
    new_full = full.replace(old, new)
    if paragraph.runs:
        paragraph.runs[0].text = new_full
        for r in paragraph.runs[1:]:
            r.text = ""
    return count


def set_paragraph_text(paragraph, new_text: str):
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(new_text)


def main():
    if not Path(INPUT).exists():
        print(f"[错误] 找不到输入文件: {INPUT}")
        return

    shutil.copy(INPUT, INPUT + ".bak")
    print(f"[备份] {INPUT}.bak")

    doc = Document(INPUT)
    log_lines = []

    def log(msg):
        print(msg)
        log_lines.append(msg)

    # ---------- 第 1 步：连续代码段折叠为占位符 ----------
    log("\n[步骤 1] 检测并替换代码块为占位符…")
    paragraphs = list(doc.paragraphs)
    i = 0
    code_block_count = 0
    while i < len(paragraphs):
        p = paragraphs[i]
        text = p.text
        if is_code_paragraph(text):
            start = i
            placeholder_name = "code_unknown.png"
            for keyword, fname in CODE_BLOCKS:
                window_start = max(0, start - 2)
                window_end = min(len(paragraphs), start + 3)
                window = "\n".join(paragraphs[j].text for j in range(window_start, window_end))
                if keyword in window or keyword in text:
                    placeholder_name = fname
                    break
            j = i
            while j < len(paragraphs) and (
                is_code_paragraph(paragraphs[j].text)
                or paragraphs[j].text.strip() == ""
                or paragraphs[j].text.strip().startswith(("#", "//"))
            ):
                j += 1
            if j - i >= 2:
                set_paragraph_text(
                    paragraphs[i],
                    f"[INSERT_IMG: {placeholder_name}]（如图所示为对应核心代码，详见仓库 app/ 目录）"
                )
                for k in range(i + 1, j):
                    set_paragraph_text(paragraphs[k], "")
                code_block_count += 1
                log(f"  · 第{i}-{j-1}段 折叠为 {placeholder_name}")
                i = j
                continue
        i += 1
    log(f"[步骤 1 完成] 共折叠 {code_block_count} 个代码块")

    # ---------- 第 2 步：全文文本替换 ----------
    log("\n[步骤 2] 执行字符串替换…")
    counters = {old: 0 for old, _ in TEXT_REPLACEMENTS}

    def walk_paragraphs(container):
        for p in container.paragraphs:
            for old, new in TEXT_REPLACEMENTS:
                c = replace_in_runs(p, old, new)
                if c:
                    counters[old] += c
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    walk_paragraphs(cell)

    walk_paragraphs(doc)
    for section in doc.sections:
        for hdr in (section.header, section.footer):
            walk_paragraphs(hdr)

    for old, c in counters.items():
        if c > 0:
            log(f"  · 「{old}」替换 {c} 次")

    # ---------- 第 3 步：图编号重排 ----------
    log("\n[步骤 3] 图编号重排…")
    fig_counters = {k: 0 for k in FIG_RENUMBER}
    for p in doc.paragraphs:
        for old, new in FIG_RENUMBER.items():
            c = replace_in_runs(p, old, new)
            if c:
                fig_counters[old] += c
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old, new in FIG_RENUMBER.items():
                        c = replace_in_runs(p, old, new)
                        if c:
                            fig_counters[old] += c
    for k, v in fig_counters.items():
        if v > 0 and not k.startswith("__TMP"):
            log(f"  · 「{k}」→「{FIG_RENUMBER[k]}」 共 {v} 次")

    # ---------- 第 4 步：缩写首次出现补全 ----------
    log("\n[步骤 4] 缩写首次出现补全…")
    done = set()
    for p in doc.paragraphs:
        for abbr, full in ABBREV_FIRST_OCCURRENCE:
            if abbr in done:
                continue
            text = p.text
            if full in text:
                done.add(abbr)
                continue
            pattern = re.compile(r'(?<![A-Za-z])' + re.escape(abbr) + r'(?![A-Za-z])')
            if pattern.search(text):
                new_text = pattern.sub(full, text, count=1)
                set_paragraph_text(p, new_text)
                done.add(abbr)
                log(f"  · 「{abbr}」→「{full}」 在第一次出现处补全")
                break
    missing = [a for a, _ in ABBREV_FIRST_OCCURRENCE if a not in done]
    if missing:
        log(f"  · 未找到首次出现位置的缩写: {missing}")

    # ---------- 保存 ----------
    doc.save(OUTPUT)
    log(f"\n[完成] 输出文件: {OUTPUT}")

    Path(LOG).write_text("\n".join(log_lines), encoding="utf-8")
    print(f"日志: {LOG}")


if __name__ == "__main__":
    main()
