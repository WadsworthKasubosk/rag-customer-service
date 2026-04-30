# -*- coding: utf-8 -*-
import docx
import io

DOCX = r'C:/Users/da983/Documents/xwechat_files/wxid_gc6r9mc9tebt22_e63d/msg/file/2026-04/正文基于RAG技术的智能客服系统的设计与实现_修订版.docx'

doc = docx.Document(DOCX)
with io.open('verified.txt', 'w', encoding='utf-8') as f:
    for i, p in enumerate(doc.paragraphs):
        f.write(f'{i:4d} [{p.style.name}] {p.text}\n')
    f.write('\n=== TABLES ===\n')
    for ti, table in enumerate(doc.tables):
        f.write(f'-- Table {ti} (rows={len(table.rows)}, cols={len(table.columns)}) --\n')
        for ri, row in enumerate(table.rows):
            cells = [c.text for c in row.cells]
            f.write(f'  row{ri}: {cells}\n')

# Count images
from docx.oxml.ns import qn
imgs = 0
for p in doc.paragraphs:
    for run in p.runs:
        for drawing in run._element.findall(qn('w:drawing')):
            imgs += 1
print(f'Inline images detected: {imgs}')
