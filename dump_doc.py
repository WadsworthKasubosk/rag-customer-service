# -*- coding: utf-8 -*-
import docx
import io
import sys

DOCX = r'C:/Users/da983/Documents/xwechat_files/wxid_gc6r9mc9tebt22_e63d/msg/file/2026-04/正文基于RAG技术的智能客服系统的设计与实现.docx'

doc = docx.Document(DOCX)

with io.open('doc_dump.txt', 'w', encoding='utf-8') as f:
    f.write('=== PARAGRAPHS ===\n')
    for i, p in enumerate(doc.paragraphs):
        f.write(f'{i:4d} [{p.style.name}] {p.text}\n')
    f.write('\n=== TABLES ===\n')
    for ti, table in enumerate(doc.tables):
        f.write(f'-- Table {ti} (rows={len(table.rows)}, cols={len(table.columns)}) --\n')
        for ri, row in enumerate(table.rows):
            cells = [c.text for c in row.cells]
            f.write(f'  row{ri}: {cells}\n')
print('done')
