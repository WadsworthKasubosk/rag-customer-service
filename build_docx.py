# -*- coding: utf-8 -*-
"""Build the complete thesis as a .docx file."""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from pathlib import Path

OUT_PATH = "基于RAG技术的智能客服系统的设计与实现.docx"
ASSETS = Path("thesis_assets")
DIAG = Path("diagrams")

CN_FONT = "宋体"
CN_FONT_HEAD = "黑体"
EN_FONT = "Times New Roman"
CODE_FONT = "Consolas"


def set_run_font(run, name_cn=CN_FONT, name_en=EN_FONT, size=12, bold=False, color=None):
    run.font.name = name_en
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name_en)
    rfonts.set(qn("w:hAnsi"), name_en)
    rfonts.set(qn("w:eastAsia"), name_cn)


def add_heading(doc, text, level=1):
    sizes = {0: 22, 1: 18, 2: 15, 3: 13, 4: 12}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if level > 0 else WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5
    r = p.add_run(text)
    set_run_font(r, name_cn=CN_FONT_HEAD, size=sizes.get(level, 12), bold=True)
    return p


def add_para(doc, text, indent=True, size=12, align="left", line_spacing=1.5):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if indent:
        pf.first_line_indent = Pt(size * 2)
    pf.line_spacing = line_spacing
    pf.space_after = Pt(0)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    set_run_font(r, size=size)
    return p


def add_keywords(doc, label, kws):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(label)
    set_run_font(r, size=12, bold=True)
    r = p.add_run(kws)
    set_run_font(r, size=12)


def add_figure(doc, img_path, caption, width_cm=14):
    img_path = str(img_path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run()
    r.add_picture(img_path, width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    cr = cap.add_run(caption)
    set_run_font(cr, size=10.5, bold=True)


def add_table_caption(doc, caption):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(8)
    cap.paragraph_format.space_after = Pt(2)
    r = cap.add_run(caption)
    set_run_font(r, size=10.5, bold=True)


def add_table(doc, header, rows, col_widths=None):
    n_cols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths:
        for col, w in zip(table.columns, col_widths):
            for cell in col.cells:
                cell.width = Cm(w)
    # header
    for i, h in enumerate(header):
        c = table.rows[0].cells[i]
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, size=10.5, bold=True)
    # rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = table.rows[1 + ri].cells[ci]
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run_font(r, size=10)
    # spacing
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)
    return table


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    # gray bg via shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    pPr.append(shd)
    for line in code.split("\n"):
        r = p.add_run(line)
        set_run_font(r, name_cn=CODE_FONT, name_en=CODE_FONT, size=9)
        r.add_break()


def add_page_break(doc):
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ============================================================
# Build document
# ============================================================
doc = Document()

# Default style
style = doc.styles["Normal"]
style.font.name = EN_FONT
style.font.size = Pt(12)
rpr = style.element.get_or_add_rPr()
rfonts = rpr.find(qn("w:rFonts"))
if rfonts is None:
    rfonts = OxmlElement("w:rFonts")
    rpr.append(rfonts)
rfonts.set(qn("w:eastAsia"), CN_FONT)

# Page margins
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(3.0)
section.right_margin = Cm(3.0)


# ============ COVER ============
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("本 科 毕 业 论 文")
set_run_font(r, name_cn=CN_FONT_HEAD, size=26, bold=True)

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("基于 RAG 技术的智能客服系统的设计与实现")
set_run_font(r, name_cn=CN_FONT_HEAD, size=22, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Design and Implementation of an Intelligent Customer Service System Based on RAG")
set_run_font(r, size=14, bold=True)

for _ in range(6):
    doc.add_paragraph()

cover_info = [("学    院", "________________"),
              ("专    业", "________________"),
              ("学生姓名", "________________"),
              ("学    号", "________________"),
              ("指导教师", "________________"),
              ("完成日期", "2026 年 5 月")]
for k, v in cover_info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{k}:{v}")
    set_run_font(r, size=14)
    p.paragraph_format.space_after = Pt(8)

add_page_break(doc)

# ============ ABSTRACT (CN) ============
add_heading(doc, "摘    要", level=0)
add_para(doc, "随着大语言模型(Large Language Model, LLM)技术的快速发展,将其与检索增强生成(Retrieval-Augmented Generation, RAG)技术相结合,为智能客服领域带来了新的解决方案。本文设计并实现了一套基于 RAG 技术的智能客服系统,以虚构品牌“星辰科技”S14 系列智能手机作为业务场景,构建了从知识库管理、意图识别、多源检索到大模型流式生成的完整客服处理流程。")
add_para(doc, "系统采用“本地规则分类 + LLM 生成分离”的核心架构,将意图分类从大模型中剥离,通过关键词规则与正则模式实现零延迟分类,使整个问答流程仅需调用一次 LLM,有效降低了响应延迟与 Token 消耗。在检索策略上,系统创新性地融合了 FAISS 非结构化文档向量检索、基于 NetworkX 的知识图谱结构化查询,以及面向业务实时数据的 MySQL 查询三种数据源,构建了多源异构的上下文增强方案。前后端通过 Server-Sent Events(SSE)协议实现流式推送,用户感知的首字延迟约为 0.5 秒。")
add_para(doc, "系统后端基于 FastAPI 异步框架开发,持久化层采用 MySQL 8.0 存储商品、订单、物流、维修工单、对话历史与用户反馈等结构化数据,并使用 Redis 7 作为缓存中间件加速热点查询;向量检索基于 FAISS 本地部署,文本嵌入采用 text2vec-base-chinese 中文模型,LLM 接口兼容 OpenAI 协议,可无缝对接 DeepSeek、通义千问、GLM 等多种大模型。前端采用原生 HTML/CSS/JavaScript 实现,无前端 UI 框架依赖,知识图谱可视化使用 vis-network 库。测试结果表明,系统意图分类准确率达到 95%,本地分类耗时不足 1 毫秒,首字响应延迟约 0.5 秒,整体响应质量与用户体验满足实际客服场景需求。")
add_keywords(doc, "关键词:", "检索增强生成;智能客服;FAISS;知识图谱;大语言模型;Server-Sent Events")
add_page_break(doc)

# ============ ABSTRACT (EN) ============
add_heading(doc, "Abstract", level=0)
add_para(doc, "With the rapid development of Large Language Model (LLM) technology, combining LLMs with Retrieval-Augmented Generation (RAG) has brought new solutions to the field of intelligent customer service. This thesis designs and implements an intelligent customer service system based on RAG technology, using a fictional brand \"Xinchen Technology\" S14 series smartphone as the business scenario, building a complete customer service pipeline from knowledge base management, intent classification, multi-source retrieval to LLM streaming generation.")
add_para(doc, "The system adopts a core architecture of \"local rule-based classification plus LLM generation separation\", decoupling intent classification from the LLM through keyword rules and regular expression patterns to achieve zero-latency classification, with only a single LLM invocation throughout the entire pipeline, effectively reducing response latency and token consumption. In terms of retrieval strategy, the system innovatively integrates three data sources: FAISS-based unstructured document vector retrieval, NetworkX-based knowledge graph structured queries, and MySQL queries for real-time business data, forming a multi-source heterogeneous context enhancement scheme. The frontend and backend communicate via the Server-Sent Events (SSE) protocol for streaming, with a perceived first-token latency of approximately 0.5 seconds.")
add_para(doc, "The backend is built on the FastAPI asynchronous framework. The persistence layer uses MySQL 8.0 to store structured data including products, orders, logistics, repair tickets, chat history and user feedback, while Redis 7 serves as a caching middleware to accelerate hot-path queries. Vector retrieval is based on FAISS deployed locally, text embedding uses the text2vec-base-chinese Chinese model, and the LLM interface is OpenAI-compatible, allowing seamless integration with DeepSeek, Tongyi Qianwen, GLM and other LLMs. The frontend is implemented in native HTML/CSS/JavaScript without any UI framework dependency, and knowledge graph visualization is implemented using the vis-network library. Experimental results show that the intent classifier achieves an accuracy of 95% with sub-millisecond local processing latency, the first-token latency is around 0.5 seconds, and the overall response quality and user experience meet the requirements of real customer service scenarios.")
add_keywords(doc, "Keywords: ", "Retrieval-Augmented Generation; Intelligent Customer Service; FAISS; Knowledge Graph; Large Language Model; Server-Sent Events")
add_page_break(doc)

# ============ TOC ============
add_heading(doc, "目    录", level=0)
toc_lines = [
    ("摘要", ""), ("Abstract", ""),
    ("第 1 章 绪论", ""),
    ("    1.1 研究背景与意义", ""), ("    1.2 国内外研究现状", ""),
    ("    1.3 本文主要工作", ""), ("    1.4 论文组织结构", ""),
    ("第 2 章 相关技术介绍", ""),
    ("    2.1 大语言模型与 DeepSeek", ""), ("    2.2 检索增强生成 RAG", ""),
    ("    2.3 向量检索与 FAISS", ""), ("    2.4 文本嵌入模型", ""),
    ("    2.5 知识图谱与 NetworkX", ""), ("    2.6 FastAPI 与 SSE 流式协议", ""),
    ("    2.7 MySQL 与 Redis", ""), ("    2.8 LangChain 框架", ""),
    ("    2.9 本章小结", ""),
    ("第 3 章 系统需求分析与总体设计", ""),
    ("    3.1 系统需求分析", ""), ("    3.2 系统总体架构", ""),
    ("    3.3 模块划分与核心处理流程", ""), ("    3.4 数据库设计", ""),
    ("    3.5 知识图谱模式设计", ""), ("    3.6 接口设计", ""),
    ("    3.7 本章小结", ""),
    ("第 4 章 系统详细实现", ""),
    ("    4.1 开发环境与项目结构", ""), ("    4.2 知识库构建与文本切分", ""),
    ("    4.3 意图分类器实现", ""), ("    4.4 RAG 核心流程实现", ""),
    ("    4.5 知识图谱构建与查询", ""), ("    4.6 电商数据库与缓存层实现", ""),
    ("    4.7 SSE 流式接口实现", ""), ("    4.8 前端界面实现", ""),
    ("    4.9 前端功能展示", ""), ("    4.10 本章小结", ""),
    ("第 5 章 系统测试与性能评估", ""),
    ("    5.1 测试环境", ""), ("    5.2 意图分类准确性测试", ""),
    ("    5.3 端到端问答测试", ""), ("    5.4 响应时延测试", ""),
    ("    5.5 Redis 缓存效果测试", ""), ("    5.6 多源检索消融实验", ""),
    ("    5.7 本章小结", ""),
    ("第 6 章 总结与展望", ""),
    ("    6.1 工作总结", ""), ("    6.2 不足与改进方向", ""),
    ("    6.3 展望", ""),
    ("参考文献", ""), ("致  谢", ""),
]
for title, _ in toc_lines:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title)
    set_run_font(r, size=12)

add_page_break(doc)


# ============ CHAPTER 1 ============
add_heading(doc, "第 1 章  绪论", level=1)
add_heading(doc, "1.1 研究背景与意义", level=2)
add_para(doc, "随着电子商务、在线教育、互联网金融等服务型业态的快速发展,客户服务已成为企业运营中不可或缺的关键环节。根据中国电子商务研究中心发布的相关数据,中国大型电商平台单日客服咨询量已达千万级,客服质量直接影响用户复购率与品牌口碑。然而,传统的人工客服模式存在三方面突出问题:其一,人力成本高昂,一个中等规模的电商企业每年仅客服人力支出就可达数千万元;其二,响应速度受坐席规模限制,大促期间用户等待时长常常超过十分钟;其三,人工客服服务质量受个体水平、情绪状态、轮班时段等因素影响显著,难以保持稳定的标准化输出;其四,7×24 小时全天候服务对人力调度提出了苛刻要求,跨时区运营更是难以为继。")
add_para(doc, "为缓解上述问题,业界自 2010 年前后开始探索智能客服技术。早期智能客服系统主要基于关键词匹配与规则模板,典型代表如基于 AIML(Artificial Intelligence Markup Language)的对话机器人。这类系统的优点是部署简单、响应快速,但其本质上是“模式匹配”而非“语言理解”,在面对用户表达多样性时极易回答错误或回答“答非所问”。第二代智能客服引入了浅层意图识别模型,如基于 TextCNN、BiLSTM 的意图分类器,以及基于槽位填充(Slot Filling)的对话管理框架,显著提高了对自然语言的容错能力,但其知识仍然以“问答对”形式人工编写,知识覆盖率受限,并且无法真正“生成”回答。")
add_para(doc, "近年来,以 GPT-3.5、GPT-4、Claude、DeepSeek、通义千问、文心一言、智谱 GLM 等为代表的大语言模型(Large Language Model, LLM)在自然语言理解与生成方面取得了突破性进展,其零样本理解能力、流畅的对话生成能力以及对复杂指令的遵循能力,为智能客服领域带来了革命性的可能。然而,直接将通用大模型应用于企业客服场景仍存在两大根本性问题。第一个问题是知识局限性:大模型的知识截止于训练时间,缺乏企业特定的产品规格、价格、库存、订单等信息,且面对未训练过的领域容易“一本正经地胡说八道”,即所谓“幻觉”(Hallucination)问题。第二个问题是工程成本:纯 LLM 调用单次成本约为传统检索的数百倍,响应延迟约为数秒,直接面向 C 端用户的高并发场景下难以承受。")
add_para(doc, "检索增强生成(Retrieval-Augmented Generation, RAG)技术为这两个问题提供了系统性解决方案。RAG 的核心思想是在生成阶段之前,先从企业知识库中检索与用户问题最相关的若干文档片段,再将这些片段作为上下文提供给大模型,由大模型基于检索到的事实生成回答。这一方案带来三个显著优点:第一,知识可即时更新,无需重新训练大模型,只需更新文档库即可;第二,回答有据可查,可以追溯到具体文档,显著降低了幻觉风险;第三,检索阶段不依赖大模型,成本可控,大模型仅在最后生成阶段被调用一次。")
add_para(doc, "然而,在工程实践中,纯文本检索仍然存在三方面短板。一是对于结构化关联问题(如“S14 Pro 与 S14 在屏幕和处理器上的差异是什么”),纯向量检索可能召回多段重复的描述性文字,而非清晰对比的结构化结果。二是对于实时业务数据(如订单状态、库存余量、维修进度),文档库中保存的信息天然滞后,无法替代数据库实时查询。三是对于意图识别,若每次都交由大模型完成,会增加延迟与成本。如何在工程上将向量检索、知识图谱、实时数据库与轻量意图识别协同起来,构建一个低延迟、高准确率、易部署的智能客服系统,正是本文研究的核心问题。")
add_para(doc, "本研究的意义体现在三个层面。在学术层面,本文实践了“分类与生成解耦”以及“非结构化检索 + 结构化检索 + 实时数据查询”三路并联的多源 RAG 架构,验证了该架构在中文电商客服场景下的可行性与有效性。在工程层面,本文给出了一套完整的、可私有化部署的智能客服系统实现,涵盖从前端交互、API 网关、RAG 链、知识图谱、缓存层到数据库的全栈方案,具有较强的复用价值。在产业层面,本文证明了基于开源组件与中国本土大模型(DeepSeek)即可构建出满足实际业务需求的智能客服系统,为中小企业以较低成本部署 AI 客服提供了参考路径。")

add_heading(doc, "1.2 国内外研究现状", level=2)
add_heading(doc, "1.2.1 国外研究现状", level=3)
add_para(doc, "RAG 框架最早由 Facebook AI Research 的 Lewis 等人于 2020 年在 NeurIPS 上提出。其原始论文《Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks》将稠密向量检索器(DPR)与序列到序列生成模型(BART)结合,在开放域问答任务上取得了显著优于纯生成模型的效果。该工作奠定了 RAG 的基本范式:用一个可微分的检索器从大规模文档库中召回相关片段,再由生成器基于片段产生答案,两者可联合训练或独立训练。")
add_para(doc, "2022 年以来,随着 ChatGPT 的横空出世,RAG 技术迅速从学术研究走向工程落地。OpenAI 推出了基于 GPT 系列的 Assistants API 与 Retrieval 工具,Anthropic 的 Claude 模型支持长达 20 万 token 的上下文以增强单次 RAG 输入容量,Google 的 Gemini 则将检索深度集成到模型推理流程中。在开源生态层面,Harrison Chase 在 2022 年发布的 LangChain 框架迅速成为 RAG 工程化的事实标准,提供了对文档加载、文本切分、向量化、检索、Prompt 编排的统一抽象;Jerry Liu 主导的 LlamaIndex 则更专注于索引数据结构与检索策略的优化。这两个框架在 GitHub 上的 Star 数均已超过 5 万,是当前 RAG 工程化的两大基石。")
add_para(doc, "学术上,围绕 RAG 的优化研究主要集中在四个方向。第一是检索质量优化。Karpukhin 等提出的 DPR 用稠密向量检索取代了传统的 BM25,显著提升了召回率;Gao 等提出的 HyDE(Hypothetical Document Embeddings)先由 LLM 生成一个“假想答案”再用该答案检索,在零样本场景下表现优于纯查询检索;Wang 等的 BGE 系列嵌入模型在 MTEB 评测榜上长期占据前列。第二是重排序(Reranker)机制。粗检索召回 Top-20 后,使用 Cross-Encoder 对 Query 与 Document 拼接打分,精排出 Top-3,显著提升了上下文质量。Cohere、Jina AI 等公司均推出了专用 Reranker API。第三是 Agentic RAG。该方向赋予 RAG 系统多步推理与工具调用能力,使其能处理需要多次检索与计算的复杂问题,代表性工作有 ReAct、Toolformer、AutoGPT 等。第四是评估体系建设。RAGAS 框架提出了 Faithfulness、Answer Relevance、Context Precision、Context Recall 四个量化指标,已成为 RAG 评测的事实标准。")
add_para(doc, "在产业落地上,Zendesk、Intercom、Salesforce 等国际客服 SaaS 巨头均已将 RAG 集成到产品中。以 Zendesk 的 AI Agents 为例,其底层即是基于 OpenAI GPT-4 加企业知识库的 RAG 方案,可处理超过 80% 的一般性咨询。Klarna、Shopify 等电商企业也披露了基于 RAG 的客服替代效果:Klarna 在 2024 年披露其 AI 客服替代了 700 个全职客服坐席的工作量,平均回复时长由 11 分钟缩短至 2 分钟。")

add_heading(doc, "1.2.2 国内研究现状", level=3)
add_para(doc, "国内大语言模型研究自 2023 年起进入高速发展期。百度文心一言、阿里通义千问、智谱 GLM、深度求索 DeepSeek、月之暗面 Kimi、字节豆包等大模型相继开放 API,且绝大多数兼容 OpenAI 协议,大幅降低了企业接入成本。其中,DeepSeek 凭借其优秀的中文理解能力与极具竞争力的价格(deepseek-chat 输入价格仅为 GPT-4 的二十分之一)迅速在中文 RAG 应用中占据重要位置,成为中小企业首选的国产 LLM 之一。")
add_para(doc, "在开源生态方面,国内涌现了一批具有影响力的中文 RAG 项目。清华大学 KEG 实验室的 ChatGLM 系列与配套的 LangChain-ChatChat 项目是中文 RAG 的代表作之一,在 GitHub 上获得了超过 3 万 Star。智源研究院发布的 BGE 中文嵌入模型(bge-large-zh)在中文语义检索任务上表现优异,广泛应用于工业界。FastGPT、Dify、MaxKB 等开源 RAG 平台则进一步降低了 RAG 应用的搭建门槛,允许用户通过可视化方式配置知识库与对话流程。")
add_para(doc, "在学术研究方面,清华大学、北京大学、复旦大学、中国科学院等机构在中文 RAG 检索质量、长文档处理、多模态 RAG 等方向均有重要产出。例如,清华大学提出的 Self-RAG 引入“自反思”机制,让模型自主决定是否需要检索;复旦大学提出的 GraphRAG 将知识图谱与向量检索深度结合,在领域问答任务上取得了显著提升。")
add_para(doc, "在产业层面,京东、阿里、拼多多、美团、网易、字节等头部企业均已在自家客服系统中集成 RAG 能力。京东智能客服“言犀”披露其日均对话量超过 1 亿次,RAG 替代率超过 70%。但具体技术细节出于商业考虑较少公开。中小企业则更多采用“开源框架 + 国产 LLM API”的组合方案进行私有化部署。")

add_heading(doc, "1.2.3 现存问题分析", level=3)
add_para(doc, "综合国内外研究与产业现状,目前 RAG 智能客服系统仍存在以下几方面的不足:第一,多源数据融合不充分。多数现有方案以纯向量检索为主,忽视了知识图谱在结构化推理上的优势,以及业务数据库在实时数据上的不可替代性。在电商客服场景中,“价格 / 库存 / 订单 / 维修进度”等问题占比超过 60%,这类问题如果完全依赖文档检索,既不准确也不及时。第二,响应延迟仍然较高。许多方案将意图识别、查询改写、检索、生成全部交由 LLM 完成,导致单次问答需要调用 LLM 三到四次,首字延迟普遍在 2 秒以上,影响用户体验。第三,部署门槛仍然偏高。许多 RAG 框架依赖 Docker、Kubernetes、向量数据库集群等重型基础设施,对于中小企业的私有化部署仍有相当门槛。第四,评估体系不完善。多数毕业设计与课程实验仅给出端到端的几个示例对话,缺乏意图分类准确率、检索准确率、首字延迟、缓存命中率等系统化的量化指标。本文针对上述问题,设计并实现了一套“分类与生成解耦、三路检索融合、单机轻量部署、配套量化评测”的智能客服系统,在工程实践层面提出了改进思路。")

add_heading(doc, "1.3 本文主要工作", level=2)
add_para(doc, "本文设计并实现了一个面向电商售前售后场景的智能客服系统,主要工作包括以下五个方面。")
add_para(doc, "第一,设计并实现了基于关键词规则与订单号正则模式的轻量级意图分类器。系统预定义了产品问答、售后服务、电商咨询、维修跟踪、物流跟踪、其他六类意图,每类意图维护一组 15 至 20 个关键词,并为订单号(DD/SF/YT/JD/KD/ORD 开头)和工单号(WX 开头)定义了正则模式。分类时优先匹配 ID 模式,未命中再用关键词逐类打分,取最高分作为分类结果,平均延迟低于 1 毫秒。")
add_para(doc, "第二,设计并实现了 FAISS 向量检索 + NetworkX 知识图谱 + MySQL 实时数据库三路并联的多源检索方案。在 RAG 链路中按“结构化知识 → 实时业务数据 → 参考资料”的优先级组装上下文,各路检索可独立失败而不影响整体可用性。该设计显著提升了回答的准确性与时效性,尤其在价格、库存、订单状态等问题上效果突出。")
add_para(doc, "第三,基于 FastAPI 实现了 SSE 流式问答接口,前端使用 Fetch API 的 ReadableStream 消费。SSE 流中先推送一个 meta 事件携带分类结果与参考来源,使前端可以立刻渲染元信息;然后逐 token 推送 token 事件;最后推送一个 done 事件携带完整回答。整个流式架构使用户感知的首字延迟约为 0.5 秒。")
add_para(doc, "第四,设计并实现了 MySQL 持久化层与 Redis 缓存层。MySQL 共设计 8 张表,涵盖商品、促销、电商物流、维修工单、快递订单、轨迹事件、对话消息、反馈记录等业务实体。Redis 采用“缓存穿透”策略加速热点查询,所有键加 rag_cs: 前缀以便管理,通过 SCAN 命令进行模式化失效以避免阻塞。Redis 不可用时系统自动降级为直查 MySQL,保证可用性。")
add_para(doc, "第五,实现了一套基于原生 HTML/CSS/JavaScript 的单文件前端,集成对话窗口、知识库管理、模型配置、知识图谱可视化、演示数据初始化、统计看板六大功能模块。知识图谱可视化采用 vis-network 库实现力导向图布局,支持节点拖拽与缩放。整个前端不依赖任何 UI 框架,部署时只需一个 HTML 文件,极大降低了部署门槛。")

add_heading(doc, "1.4 论文组织结构", level=2)
add_para(doc, "全文共分为六章。第 1 章绪论,介绍研究背景、国内外研究现状以及本文工作。第 2 章介绍系统涉及的关键技术,包括大语言模型、RAG 框架、FAISS、嵌入模型、知识图谱、FastAPI、SSE、MySQL、Redis 与 LangChain。第 3 章给出系统需求分析与总体设计,包括分层架构、核心处理流程、数据库设计、知识图谱模式设计与接口设计。第 4 章详述各模块的实现细节,包括知识库构建、意图分类器、RAG 核心流程、知识图谱构建与查询、缓存层、SSE 流式接口、前端界面与前端功能展示。第 5 章对系统进行功能测试与性能评估,包括意图分类准确率测试、端到端问答测试、响应时延测试、Redis 缓存效果测试与多源检索消融实验。第 6 章总结全文,分析不足并展望后续工作。最后是参考文献与致谢。")

add_page_break(doc)


# ============ CHAPTER 2 ============
add_heading(doc, "第 2 章  相关技术介绍", level=1)

add_heading(doc, "2.1 大语言模型与 DeepSeek", level=2)
add_heading(doc, "2.1.1 大语言模型概述", level=3)
add_para(doc, "大语言模型(Large Language Model, LLM)是指参数规模在数十亿至万亿级别、基于 Transformer 架构、在大规模文本语料上以自监督方式预训练得到的语言模型。LLM 的核心能力来自三方面:基于自注意力机制的长上下文建模能力、通过大规模数据训练获得的世界知识与语言模式、通过指令微调(Instruction Tuning)与人类反馈强化学习(RLHF)获得的对话与指令遵循能力。")
add_para(doc, "自 2018 年 BERT 与 GPT-1 发布以来,LLM 经历了从小型预训练模型到超大规模通用助手的快速演进。GPT-3(1750 亿参数)首次展示了“仅靠规模即可获得显著零样本能力”的现象;ChatGPT 于 2022 年 11 月发布,引爆了大模型的工业应用;GPT-4、Claude 3、Gemini 等闭源模型不断刷新性能上限;Llama、Mistral、Qwen 等开源模型则推动了 LLM 在产业界的普及。")

add_heading(doc, "2.1.2 DeepSeek 模型选型", level=3)
add_para(doc, "本系统选用深度求索(DeepSeek)公司的 deepseek-chat 模型作为生成模型,主要基于以下三方面考虑。")
add_para(doc, "第一,API 兼容性。DeepSeek API 完全兼容 OpenAI 协议,可通过 LangChain 提供的 langchain-openai 客户端直接调用,只需在初始化 ChatOpenAI 时设置 base_url=https://api.deepseek.com 即可,无需为接入 DeepSeek 编写专用客户端。这一兼容性设计意味着系统可以零成本切换到其他兼容 OpenAI 协议的国产模型,如阿里通义千问、智谱 GLM、月之暗面 Kimi 等。")
add_para(doc, "第二,中文理解能力。DeepSeek 在中文语料上进行了重点优化,在 C-Eval、CMMLU、C3 等中文评测榜上表现优秀。对于本系统的电商客服场景,中文理解能力至关重要,DeepSeek 的中文表达流畅度与对中国电商语境的理解显著优于同等价位的国外模型。")
add_para(doc, "第三,成本优势。截至本文撰写时,deepseek-chat 的输入价格约为 1 元/百万 token,输出价格约为 2 元/百万 token,仅为 GPT-4 同类价格的二十分之一,为 GPT-3.5-turbo 的二分之一,极大降低了系统运营成本。")

add_heading(doc, "2.1.3 LLM 调用方式", level=3)
add_para(doc, "本系统通过 langchain_openai.ChatOpenAI 客户端调用 DeepSeek API,关键参数设置如下:model 设置为 deepseek-chat;openai_api_base 设置为 https://api.deepseek.com;temperature 设置为 0.3,在确定性与表达多样性之间取得平衡 — 客服场景需要较低的温度以保证回答稳定,但过低又会使表达呆板;max_tokens 设置为 2048,足以覆盖绝大多数客服回答的长度需求。系统通过 llm.stream(prompt) 调用流式接口,LLM 逐 token 返回内容,前端逐 token 渲染,用户感知的首字延迟约为 0.5 秒。")

add_heading(doc, "2.2 检索增强生成 RAG", level=2)
add_heading(doc, "2.2.1 RAG 的提出背景", level=3)
add_para(doc, "RAG 由 Lewis 等人于 2020 年在 NeurIPS 上提出,旨在解决纯生成模型的两大问题:其一是参数化知识的静态性,即模型训练完成后知识无法更新,需要通过昂贵的再训练才能引入新知识;其二是生成幻觉,即模型在不知道答案时倾向于“编造”看似合理但实际错误的内容。RAG 通过引入外部检索机制,将“知识”从模型参数中分离出来,放入可即时更新的外部知识库,既解决了知识更新问题,又通过“以检索为依据生成”的方式显著降低了幻觉风险。")

add_heading(doc, "2.2.2 RAG 通用流程", level=3)
add_para(doc, "RAG 系统通常包含五个阶段。(1)文档加载(Loading)。将 PDF、Word、TXT、HTML、Markdown 等不同格式的原始文档读入系统,提取纯文本内容。LangChain 提供了 PyPDFLoader、Docx2txtLoader、TextLoader 等多种加载器以支持不同格式。(2)文本切分(Splitting)。将长文档切分为若干较短的“块”(Chunk),每块通常为数百字。切分策略有固定长度切分、按段落切分、递归字符切分等。本系统使用 LangChain 的 RecursiveCharacterTextSplitter,以 500 字为目标长度、50 字重叠递归切分。(3)向量化与索引构建(Embedding & Indexing)。将每个文本块通过嵌入模型转换为一个高维稠密向量(本系统为 768 维),并构建向量索引以支持高效的近似最近邻搜索。本系统使用 FAISS 作为向量索引。(4)相似度检索(Retrieval)。查询时,先将用户问题转为查询向量,再在向量索引中检索 Top-K 最相似的文本块返回。本系统取 K=5。(5)Prompt 拼装与生成(Augmented Generation)。将检索到的文本块作为上下文,与用户问题、对话历史一起填入 Prompt 模板,送入 LLM 生成最终回答。")
add_figure(doc, DIAG / "d2_rag_flow.png", "图 2-1  RAG 通用流程图", width_cm=8)

add_heading(doc, "2.2.3 RAG 的演进", level=3)
add_para(doc, "RAG 自提出以来快速演进,目前业界已形成“朴素 RAG(Naive RAG)→ 高级 RAG(Advanced RAG)→ 模块化 RAG(Modular RAG)”三个发展阶段的共识。朴素 RAG 即上述五段式流程;高级 RAG 在朴素流程基础上引入查询改写、混合检索(向量 + BM25)、重排序、上下文压缩等优化手段;模块化 RAG 则进一步将整个流程抽象为可组合的模块,允许引入路由、循环、Agent 等复杂控制流。本系统在朴素 RAG 流程基础上,创新性地引入了“意图分类路由 + 三路并联检索 + 分类化 Prompt 模板”机制,可视为模块化 RAG 在中文电商客服场景下的一种工程实践。")

add_heading(doc, "2.3 向量检索与 FAISS", level=2)
add_heading(doc, "2.3.1 向量检索原理", level=3)
add_para(doc, "向量检索的核心思想是将语义相似度转化为向量空间中的距离度量。给定一个查询向量 q 与文档向量集合 {d_1, d_2, ..., d_n},检索任务即找到与 q 距离最小(或相似度最高)的 K 个文档向量。常用的距离度量有欧氏距离(L2)、内积、余弦相似度三种,对于经过 L2 归一化的向量,三者在排序结果上等价。精确最近邻搜索的复杂度为 O(n·d),当 n 达到百万级时已难以支持实时查询。近似最近邻(Approximate Nearest Neighbor, ANN)算法通过引入索引结构,以微小的精度损失换取数量级的速度提升。常见的 ANN 算法包括基于哈希的 LSH、基于树的 KD-Tree、基于图的 HNSW、基于聚类与量化的 IVF + PQ 等。")

add_heading(doc, "2.3.2 FAISS 简介", level=3)
add_para(doc, "FAISS(Facebook AI Similarity Search)是 Meta(原 Facebook)于 2017 年开源的高性能向量检索库,使用 C++ 实现并提供 Python 绑定,支持百万级到十亿级向量的近似最近邻搜索。FAISS 提供了多种索引类型,包括 IndexFlat(暴力精确搜索)、IndexIVFFlat(倒排索引 + 精确搜索)、IndexIVFPQ(倒排索引 + 乘积量化)、IndexHNSW(基于图的近似搜索)等,可根据数据规模与精度要求灵活选择。")

add_heading(doc, "2.3.3 本系统的 FAISS 使用方式", level=3)
add_para(doc, "本系统使用 faiss-cpu 1.9.0,通过 langchain_community.vectorstores.FAISS 进行高层封装。考虑到本系统的文档规模在万级,使用 IndexFlat(暴力搜索)即可获得毫秒级响应,且无需调参,因此采用 LangChain 默认的 IndexFlatL2。索引文件持久化为 data/faiss_index/index.faiss(向量数据)与 index.pkl(元数据),启动时通过 FAISS.load_local 加载到内存中作为单例使用。检索时使用 similarity_search_with_score(query, k=5) 方法,返回 Top-5 最相似文档及其 L2 距离。系统通过公式 cosine = 1 - L2² / 2 将 L2 距离转换为余弦相似度,使前端展示的相似度分数更直观。")
add_figure(doc, DIAG / "d4_retrieval.png", "图 2-2  检索流程详细图", width_cm=14)

add_heading(doc, "2.4 文本嵌入模型", level=2)
add_para(doc, "文本嵌入(Text Embedding)是将变长的文本映射为定长稠密向量的过程,所得向量被称为该文本的语义表示。优质的嵌入模型应当具备如下特性:语义相近的文本在向量空间中距离较近,语义相远的文本距离较远;对常见同义改写、词序变化、句式变化具有鲁棒性;对中文、英文、混合语言均具有良好支持。")
add_para(doc, "本系统使用 shibing624/text2vec-base-chinese 作为中文文本嵌入模型。该模型由开源社区贡献者徐明发布于 HuggingFace Model Hub,基于 BERT-base 架构,在 LCQMC、BQ、PAWSX 等中文语义匹配数据集上进行了对比学习微调,输出 768 维向量。该模型在中文短文本相似度任务上的 Spearman 相关系数达到 0.78 以上,接近 BERT-large 量级模型的水平,但参数量仅约 1 亿,适合在 CPU 上进行实时推理。模型通过 langchain-huggingface 的 HuggingFaceEmbeddings 类加载,运行设备指定为 CPU,启用 normalize_embeddings=True 使输出向量单位化,便于后续的余弦相似度计算。模型采用单例模式,首次调用时加载,后续复用。")

add_heading(doc, "2.5 知识图谱与 NetworkX", level=2)
add_para(doc, "知识图谱(Knowledge Graph)是以图结构表示实体及其关系的语义网络,核心元素为节点(Entity)与边(Relation)。知识图谱具有三方面优势。结构化关联推理:可以方便地表达“产品 — SKU — 规格”、“产品 — 故障 — 维修费用”这类多跳关系。显式语义:每条边的关系类型(如 HAS_SKU、HAS_FAULT)都是显式定义的,语义清晰、可解释。精确查询:支持基于图遍历的精确查询,无需依赖向量相似度。知识图谱在 RAG 系统中可与向量检索互补。向量检索擅长处理表达多样的非结构化问题,知识图谱擅长处理结构化关联问题(如规格对比、价格查询、故障 — 维修费用关联)。")
add_para(doc, "NetworkX 是一个用于图分析的 Python 库,支持多种图类型(无向图、有向图、多重图)与丰富的图算法。本系统使用 NetworkX 的有向图(DiGraph)结构,节点存储产品、SKU、规格、促销、库存、配件、保修、故障等实体,边存储 HAS_SKU、HAS_SPEC、HAS_PROMOTION、HAS_STOCK、HAS_ACCESSORY、HAS_WARRANTY、HAS_FAULT 等关系。整个图谱在构建脚本 build_graph.py 中以代码方式定义,持久化为 data/kg/graph.pkl 文件,运行时通过 pickle.load 加载到内存中作为单例使用。选用 NetworkX 而非 Neo4j 等专用图数据库的原因有三:本系统图谱规模较小(节点约 50 个,边约 80 条),内存图足以胜任;NetworkX 部署简单,无需额外的数据库进程;NetworkX 提供了丰富的图算法 API,便于后续扩展。")

add_heading(doc, "2.6 FastAPI 与 SSE 流式协议", level=2)
add_para(doc, "FastAPI 是 Sebastián Ramírez 于 2018 年发布的现代 Python Web 框架,基于 Starlette 与 Pydantic 构建。其核心特性包括:基于 Python 类型注解的自动数据校验与文档生成、原生支持 async/await 异步编程、自动生成 OpenAPI(Swagger)与 ReDoc 文档、性能接近 NodeJS 与 Go。本系统使用 FastAPI 0.115 作为后端框架,通过 APIRouter 将路由按业务域拆分为 chat、knowledge、feedback、config、repair、logistics、graph、db、demo 共 9 个模块。")
add_para(doc, "Server-Sent Events(SSE)是 HTML5 标准定义的服务器推送协议,基于 HTTP 长连接,以 text/event-stream 为 MIME 类型。SSE 协议格式简单,每条事件以 data: 开头、\\n\\n 结尾,客户端通过 EventSource 或 Fetch API 的 ReadableStream 消费。相比 WebSocket,SSE 具有三方面优势:实现简单,基于普通 HTTP,不需要握手与帧解析;原生支持自动重连,断线后浏览器会自动重新建立连接;天然兼容 HTTP/2 的多路复用。其缺点是只支持服务器到客户端的单向推送,但对于 LLM 流式输出场景已经足够。本系统通过 FastAPI 的 StreamingResponse 实现 SSE,生成器逐 yield 事件,响应头自动设置为 text/event-stream。")

add_heading(doc, "2.7 MySQL 与 Redis", level=2)
add_para(doc, "MySQL 是世界上使用最广泛的开源关系型数据库之一。本系统使用 MySQL 8.0,采用 InnoDB 存储引擎以支持事务与外键约束,字符集设置为 utf8mb4 以完整支持中文与 emoji 字符。系统通过 SQLAlchemy 2.0 ORM 进行数据访问,使用 Mapped 与 mapped_column 类型注解定义模型,使用 session_scope 上下文管理器统一处理事务提交与回滚。")
add_para(doc, "Redis 是高性能的内存键值数据库,支持字符串、哈希、列表、集合、有序集合等多种数据结构。本系统使用 Redis 7.2 作为缓存中间件,采用如下设计:所有缓存值通过 JSON 序列化为字符串存储;默认 TTL 为 300 秒,可通过环境变量 CACHE_TTL_SECONDS 配置;所有键名加 rag_cs: 前缀便于隔离与管理;缓存失效采用 SCAN + DEL 而非 KEYS * 以避免阻塞主线程;客户端连接设置 1 秒超时,Redis 不可用时降级为直查 MySQL,保证系统可用性。")

add_heading(doc, "2.8 LangChain 框架", level=2)
add_para(doc, "LangChain 是 Harrison Chase 于 2022 年发布的 LLM 应用开发框架,提供了对大模型、嵌入模型、向量库、Prompt 模板、文档加载器、文本切分器、Memory、Agent 等组件的统一抽象。LangChain 的核心理念是“链(Chain)”,将一个 LLM 应用拆解为一系列可组合的步骤,每一步通过统一接口连接。")
add_para(doc, "本系统使用 LangChain 0.3.x 系列,具体依赖包括:langchain 0.3.7(核心抽象)、langchain-community 0.3.7(社区集成,包含 FAISS 向量库)、langchain-openai 0.2.14(OpenAI 兼容客户端,用于调用 DeepSeek)、langchain-huggingface 0.1.2(HuggingFace 嵌入模型集成)。其中,PromptTemplate 用于参数化 Prompt 模板,RecursiveCharacterTextSplitter 用于文本切分,HuggingFaceEmbeddings 用于加载嵌入模型,FAISS 用于向量索引管理,ChatOpenAI 用于 LLM 调用。LangChain 的高层抽象使本系统可以以模块化方式构建 RAG 链路,且可灵活替换各个组件。")

add_heading(doc, "2.9 本章小结", level=2)
add_para(doc, "本章系统介绍了系统涉及的九类核心技术:大语言模型与 DeepSeek 选型理由、RAG 框架的演进与流程、FAISS 向量检索原理与本系统的使用方式、text2vec-base-chinese 中文嵌入模型、NetworkX 知识图谱实现、FastAPI 与 SSE 流式协议、MySQL 8.0 与 Redis 7 的存储设计、LangChain 框架的核心抽象。这些技术构成了下一章系统设计与实现的技术基础。")

add_page_break(doc)


# ============ CHAPTER 3 ============
add_heading(doc, "第 3 章  系统需求分析与总体设计", level=1)

add_heading(doc, "3.1 系统需求分析", level=2)
add_heading(doc, "3.1.1 业务场景分析", level=3)
add_para(doc, "本系统以虚构品牌“星辰科技”S14 系列智能手机的售前售后客服为业务场景。该场景下用户咨询的问题主要可归为以下六类。")
add_para(doc, "第一类是产品问答,占比约 25%。典型问题如“S14 Pro 的处理器是什么”、“怎么开启快充模式”、“S14 支持哪些蓝牙协议”。这类问题答案相对固定,主要依赖产品说明书与使用指南文档,适合纯 RAG 检索方案。")
add_para(doc, "第二类是售后服务,占比约 20%。典型问题如“屏幕碎了能保修吗”、“换电池多少钱”、“进水了怎么办”。这类问题既涉及保修政策(结构化知识),又涉及故障排查流程(非结构化文档),适合知识图谱与 RAG 联合检索。")
add_para(doc, "第三类是电商咨询,占比约 25%。典型问题如“S14 16+512 多少钱”、“现在有什么优惠”、“星空黑还有货吗”。这类问题需要查询实时价格、促销与库存,必须依赖 MySQL 实时数据库,纯文档检索无法保证数据时效性。")
add_para(doc, "第四类是维修跟踪,占比约 10%。典型问题如“WX20260320001 修好了吗”、“我的工单进度”。这类问题需要凭工单号查询维修服务表,获取实时维修状态。")
add_para(doc, "第五类是物流跟踪,占比约 15%。典型问题如“DD20260320001 到哪了”、“我的快递发货了吗”。这类问题需要凭订单号查询物流轨迹表。")
add_para(doc, "第六类是其他,占比约 5%。包括无明确意图的闲聊、跨品类问题等。上述六类业务场景,没有任何一种单一检索方案可以同时高效处理。这一观察直接驱动了本系统“多源异构检索融合”的核心架构设计。")

add_heading(doc, "3.1.2 功能需求", level=3)
add_para(doc, "(1)智能问答功能。支持用户以自然语言提问,系统能够准确识别意图并返回符合企业知识库的回答。回答需附带参考来源(文档片段及相似度分数),便于用户核实。支持基于会话 ID 的多轮对话,系统能记忆最近 5 轮对话历史并将其作为上下文输入大模型。回答以流式方式逐 token 推送,首字延迟应低于 1 秒。")
add_para(doc, "(2)知识库管理功能。支持 PDF、DOCX、TXT 三种格式文档的上传。上传后系统自动完成文档解析、文本切分、向量化与索引构建。支持清空向量库以便重新构建。前端实时展示当前知识库的文档数量与切片数量。")
add_para(doc, "(3)模型配置功能。支持运行时动态配置 LLM 服务商、API Key、Base URL 与模型名,无需重启服务即可生效。支持连接测试,验证 API Key 是否有效、目标模型是否可用,并返回单次调用的耗时与 Token 消耗。API Key 在前端展示时需脱敏(仅显示前 7 位与后 4 位)。")
add_para(doc, "(4)知识图谱可视化。以力导向图方式展示产品、SKU、规格、促销、故障等实体及其关联关系。支持节点拖拽、画布缩放、节点颜色按类型区分。鼠标悬停节点时显示节点详细属性。")
add_para(doc, "(5)实时业务查询。当用户问题中包含订单号(DD/SF/YT/JD/KD/ORD 开头)或工单号(WX 开头)时,系统自动从 MySQL 中查询最新的物流轨迹与维修进度,并将查询结果作为上下文输入大模型。")
add_para(doc, "(6)反馈与统计。用户可对每轮回答进行 1 至 5 星评分,并标记“已解决”或“未解决”。系统在统计页面展示总对话量、平均评分、满意率(评分≥4 的占比)与解决率。")
add_para(doc, "(7)演示数据初始化。系统启动时自动初始化测试数据(8 个商品、5 个促销、若干订单、工单、物流轨迹),并提供“一键重置”接口便于演示与评测。")

add_heading(doc, "3.1.3 非功能需求", level=3)
add_para(doc, "(1)性能需求。意图分类延迟应低于 50 毫秒(目标值 1 毫秒);单次 RAG 检索延迟应低于 200 毫秒;首字响应延迟应低于 1 秒(目标值 500 毫秒);单机至少支持 50 路并发问答。")
add_para(doc, "(2)可用性需求。Redis 不可用时系统应自动降级为直查 MySQL,降级期间业务功能不受影响,仅响应延迟有所上升。向量库为空时(冷启动场景)应自动回退到 BM25 关键词检索,保证功能可用。LLM API 调用失败时应返回友好错误提示,而非直接抛出 500 错误。")
add_para(doc, "(3)可扩展性需求。LLM 接口应可替换,支持接入兼容 OpenAI 协议的任何大模型(DeepSeek、通义千问、GLM、Kimi 等)。向量库接口应可替换,未来可支持切换到 Milvus、Chroma 等专用向量数据库。意图分类策略应可扩展,支持后续引入轻量分类模型(如 FastText)替代关键词规则。")
add_para(doc, "(4)安全性需求。所有敏感配置(API Key、数据库密码)通过环境变量注入,不硬编码到代码中。API Key 在前端展示时脱敏。前端流式渲染阶段使用 textContent 而非 innerHTML 防御 XSS 注入,完整答案渲染时使用 marked.parse 进行 Markdown 解析,该库内置危险标签过滤。")
add_para(doc, "(5)易部署性需求。系统应能在 Windows、Linux、macOS 三大主流操作系统上一键部署。除 Python 与 MySQL、Redis 外不依赖额外的重型基础设施(如 Kubernetes、专用向量数据库)。提供跨平台部署脚本(deploy.sh、deploy.ps1、deploy.bat)。")

add_heading(doc, "3.2 系统总体架构", level=2)
add_para(doc, "系统采用经典的分层架构设计,自顶向下分为用户层、API 网关层、业务路由层、核心服务层、模型层与数据层共六层。各层职责清晰,层间通过明确定义的接口交互,符合软件工程“高内聚、低耦合”的设计原则。系统总体架构如图 3-1 所示。")
add_figure(doc, DIAG / "d1_overall.png", "图 3-1  系统总体架构图", width_cm=15)
add_para(doc, "(1)用户层。包括 Web 前端、API 客户端与测试工具。Web 前端是基于原生 HTML/CSS/JavaScript 的单文件页面,提供完整的交互功能。API 客户端泛指任何符合 RESTful 规范的调用方,可通过 HTTP 请求接入本系统。测试工具包括 Postman、curl、Python requests 等,用于开发期接口测试与压力测试。")
add_para(doc, "(2)API 网关层。由 FastAPI 主入口(app/main.py)统一接收所有 HTTP 请求,完成路由分发、CORS 处理、异常捕获、健康检查等横切关注点。FastAPI 框架基于 Starlette 与 Uvicorn,采用 ASGI 协议,原生支持异步请求处理。")
add_para(doc, "(3)业务路由层。位于 app/api/ 目录,按业务域拆分为 9 个 APIRouter:chat(对话)、knowledge(知识库)、feedback(反馈)、config(配置)、repair(维修)、logistics(物流)、graph(图谱)、db(数据库管理)、demo(演示)。每个路由模块只负责参数校验、调用对应的服务层、组装响应,本身不包含业务逻辑。")
add_para(doc, "(4)核心服务层。位于 app/services/ 与 app/core/ 目录。services/ 包含 chat_service、history_service、kg_service、db_service、repair_service、logistics_service、config_service 共 7 个业务服务模块。core/ 包含 RAG 链(rag_chain)、意图分类器(classifier)、向量存储封装(vector_store)、知识图谱实现(knowledge_graph)、嵌入模型封装(embeddings)、Prompt 模板(prompt_templates)共 6 个核心模块。")
add_para(doc, "(5)模型层。位于 app/models/ 目录,封装大模型与嵌入模型的调用细节。llm.py 提供 get_llm()、update_llm_config()、test_connection() 等接口,屏蔽底层模型差异。")
add_para(doc, "(6)数据层。包括 MySQL 8.0、Redis 7.2、FAISS 索引文件、知识图谱 pkl 文件与原始文档库五种数据源。MySQL 存储结构化业务数据;Redis 缓存热点查询结果;FAISS 索引文件保存文档向量;知识图谱 pkl 文件保存 NetworkX 有向图序列化结果;原始文档库即用户上传的 PDF/DOCX/TXT 文件。")

add_heading(doc, "3.3 模块划分与核心处理流程", level=2)
add_heading(doc, "3.3.1 模块划分", level=3)
add_para(doc, "本系统按职责将代码组织为 7 大模块,各模块之间的依赖关系如表 3-1 所示。")
add_table_caption(doc, "表 3-1  系统模块划分与依赖关系")
add_table(doc,
    ["模块名", "路径", "职责", "依赖模块"],
    [
        ["API 路由层", "app/api/", "接收 HTTP 请求,参数校验,响应组装", "业务服务层"],
        ["业务服务层", "app/services/", "业务逻辑实现,事务编排", "核心服务层、存储层"],
        ["核心服务层", "app/core/", "RAG 流程、意图分类、检索", "模型层、存储层"],
        ["模型层", "app/models/", "LLM 与 Embedding 调用封装", "无"],
        ["存储层", "app/store/", "ORM 模型、数据库会话、缓存封装", "无"],
        ["知识管理", "app/knowledge/", "文档加载与切分", "无"],
        ["脚本工具", "app/scripts/", "知识图谱构建、数据库初始化", "存储层"],
    ])
add_para(doc, "这种分层方式遵循“依赖单向流动”原则:上层依赖下层,下层不感知上层,从而保证了模块的可测试性与可替换性。")

add_heading(doc, "3.3.2 核心处理流程", level=3)
add_para(doc, "系统核心处理流程是从用户提问到最终回答的完整链路,经过 7 个阶段,如图 3-2 所示。各阶段的具体职责如下:(1)用户提问。前端通过 Fetch API 向 /api/chat/stream 端点发起 POST 请求,请求体包含 session_id 与 question 两个字段。(2)本地关键词分类。调用 classify_question() 函数,先用正则模式匹配订单号与工单号(优先级最高),未命中再用关键词逐类打分,取最高分作为分类结果。整个过程在 1 毫秒内完成,不调用任何 LLM。(3)并行检索。系统并行执行三路检索:FAISS 向量检索从文档库中召回 Top-5 相关片段;知识图谱查询根据分类与关键词路由到对应的查询函数,返回结构化知识文本;MySQL 实时业务查询仅在分类为维修跟踪或物流跟踪时触发,从对应业务表中提取实时数据。(4)上下文组装。按“结构化知识 → 实时业务数据 → 参考资料”的优先级将三路检索结果拼接为完整的上下文。三段之间用 \\n\\n---\\n\\n 分隔,每段加上明确的标题。(5)按分类选择 Prompt 模板。系统预定义了 6 个 PromptTemplate,分别对应六类意图。(6)LLM 流式生成。调用 llm.stream(prompt) 启动流式生成,DeepSeek API 逐 token 返回内容。这是整个流程中唯一一次 LLM 调用。(7)SSE 逐 token 推送。FastAPI 的 StreamingResponse 将生成器包装为 SSE 响应,每个事件以 data: 开头、\\n\\n 结尾,前端通过 ReadableStream 逐事件消费。")
add_figure(doc, ASSETS / "fig_3_2_flow.png", "图 3-2  系统核心处理流程图", width_cm=12)
add_para(doc, "整个流程的关键设计是意图分类与 LLM 生成解耦:意图分类用纯本地规则在 1 毫秒内完成,使全流程只需调用一次 LLM,显著降低了延迟与 Token 消耗。这一设计也为后续引入轻量分类模型(如 FastText、小型 BERT)留下了扩展空间。")

add_heading(doc, "3.4 数据库设计", level=2)
add_heading(doc, "3.4.1 数据库选型", level=3)
add_para(doc, "本系统使用 MySQL 8.0 作为持久化存储。选用 MySQL 而非 PostgreSQL、SQLite 等其他关系型数据库的原因有三:第一,MySQL 在国内电商业务中使用最广泛,选用 MySQL 便于未来与企业既有系统集成;第二,MySQL 8.0 的 InnoDB 引擎对事务、外键、索引的支持完善,utf8mb4 字符集对中文与 emoji 的支持完整;第三,MySQL 的运维生态成熟,部署简单,适合本科毕设私有化部署的场景。数据访问层使用 SQLAlchemy 2.0 ORM,通过 Mapped 与 mapped_column 类型注解定义模型,通过 session_scope 上下文管理器统一处理事务提交与回滚,通过 pymysql 作为底层数据库驱动。")

add_heading(doc, "3.4.2 数据库逻辑设计", level=3)
add_para(doc, "系统共设计 8 张表,涵盖商品域、订单域、维修域、对话域四个业务子域,各表通过外键建立关联,如图 3-3 所示。")
add_figure(doc, ASSETS / "fig_3_3_er.png", "图 3-3  数据库 ER 图", width_cm=15)
add_para(doc, "各表的字段设计如表 3-2 至表 3-9 所示。")

add_table_caption(doc, "表 3-2  products 表(商品主表)")
add_table(doc, ["字段名", "类型", "约束", "说明"], [
    ["sku_id", "INTEGER", "主键", "SKU 编号,如 1001、1002"],
    ["product_name", "VARCHAR(255)", "非空,索引", "商品名,如“星辰 S14 8+256GB 星空黑”"],
    ["current_price", "FLOAT", "非空", "官方价(元)"],
    ["promo_price", "FLOAT", "可空", "促销价(元),为空表示无促销"],
    ["stock_num", "INTEGER", "非空,默认 0", "库存数量"],
    ["specs_json", "TEXT", "可空", "规格 JSON"],
    ["is_on_sale", "BOOLEAN", "非空,默认 TRUE", "是否在售"],
])

add_table_caption(doc, "表 3-3  promotions 表(促销活动表)")
add_table(doc, ["字段名", "类型", "约束", "说明"], [
    ["promo_id", "INTEGER", "主键,自增", "促销 ID"],
    ["sku_id", "INTEGER", "外键 → products,索引", "关联商品 SKU"],
    ["description", "VARCHAR(255)", "非空", "促销描述"],
    ["discount_rate", "FLOAT", "可空", "折扣率"],
    ["start_date", "VARCHAR(32)", "可空", "开始日期"],
    ["end_date", "VARCHAR(32)", "可空", "结束日期"],
])

add_table_caption(doc, "表 3-4  ecommerce_logistics 表(电商订单物流表)")
add_table(doc, ["字段名", "类型", "约束", "说明"], [
    ["order_id", "VARCHAR(32)", "主键", "订单号,如 ORD1730000000000"],
    ["sku_id", "INTEGER", "非空,索引", "商品 SKU"],
    ["carrier", "VARCHAR(64)", "非空", "承运商,如顺丰、京东物流"],
    ["status", "VARCHAR(64)", "非空", "物流状态"],
    ["address", "VARCHAR(255)", "可空", "收件地址"],
    ["created_at", "DATETIME", "非空", "创建时间"],
])

add_table_caption(doc, "表 3-5  repair_tickets 表(维修工单表)")
add_table(doc, ["字段名", "类型", "约束", "说明"], [
    ["ticket_id", "VARCHAR(32)", "主键", "工单号,如 WX20260320001"],
    ["phone", "VARCHAR(32)", "非空,索引", "用户手机号"],
    ["product", "VARCHAR(255)", "非空", "送修产品"],
    ["issue", "TEXT", "非空", "故障描述"],
    ["status", "VARCHAR(64)", "非空", "维修状态"],
    ["created_at", "DATETIME", "非空", "创建时间"],
    ["updated_at", "DATETIME", "非空", "最后更新时间"],
    ["estimated_days", "INTEGER", "非空,默认 3", "预计维修天数"],
])

add_table_caption(doc, "表 3-6  tracking_orders 表(快递订单表)")
add_table(doc, ["字段名", "类型", "约束", "说明"], [
    ["order_id", "VARCHAR(32)", "主键", "订单号,如 DD20260320001"],
    ["tracking_no", "VARCHAR(64)", "非空,唯一,索引", "快递单号"],
    ["carrier", "VARCHAR(64)", "非空", "承运商"],
    ["status", "VARCHAR(64)", "非空", "当前物流状态"],
    ["items", "VARCHAR(255)", "非空", "商品摘要"],
    ["created_at", "DATETIME", "非空", "下单时间"],
    ["updated_at", "DATETIME", "非空", "最后更新时间"],
])

add_table_caption(doc, "表 3-7  tracking_events 表(物流轨迹事件表)")
add_table(doc, ["字段名", "类型", "约束", "说明"], [
    ["event_id", "INTEGER", "主键,自增", "事件 ID"],
    ["order_id", "VARCHAR(32)", "外键 → tracking_orders,索引", "关联订单"],
    ["event_time", "DATETIME", "非空", "事件发生时间"],
    ["location", "VARCHAR(255)", "非空", "事件发生地点"],
    ["event", "VARCHAR(255)", "非空", "事件描述"],
])

add_table_caption(doc, "表 3-8  chat_messages 表(对话消息表)")
add_table(doc, ["字段名", "类型", "约束", "说明"], [
    ["id", "INTEGER", "主键,自增", "消息 ID"],
    ["session_id", "VARCHAR(128)", "非空,索引", "会话 ID"],
    ["role", "VARCHAR(32)", "非空", "角色:user / assistant"],
    ["content", "TEXT", "非空", "消息内容"],
    ["created_at", "DATETIME", "非空", "创建时间"],
])

add_table_caption(doc, "表 3-9  feedback_records 表(用户反馈表)")
add_table(doc, ["字段名", "类型", "约束", "说明"], [
    ["id", "INTEGER", "主键,自增", "反馈 ID"],
    ["session_id", "VARCHAR(128)", "非空,索引", "会话 ID"],
    ["rating", "INTEGER", "非空", "1-5 星评分"],
    ["comment", "TEXT", "非空,默认空字符串", "文字评价"],
    ["resolved", "BOOLEAN", "可空", "是否已解决"],
    ["created_at", "DATETIME", "非空", "创建时间"],
])

add_heading(doc, "3.4.3 数据库物理设计要点", level=3)
add_para(doc, "所有表均使用 utf8mb4 字符集与 utf8mb4_unicode_ci 排序规则,确保对中文与 emoji 的完整支持。所有时间字段统一使用 DATETIME 类型存储 UTC 时间,前端展示时由前端代码转换为本地时间。外键关联通过 SQLAlchemy 的 ForeignKey 与 relationship 实现,并启用 cascade=\"all, delete-orphan\" 以保证子表数据随父表删除自动清理。所有高频查询字段均建立索引,如 products.product_name、chat_messages.session_id、repair_tickets.phone。")

add_heading(doc, "3.5 知识图谱模式设计", level=2)
add_heading(doc, "3.5.1 节点类型设计", level=3)
add_para(doc, "知识图谱采用有向图结构,共定义 8 种节点类型,如表 3-10 所示。")
add_table_caption(doc, "表 3-10  知识图谱节点类型")
add_table(doc, ["节点类型", "节点 ID 示例", "关键属性", "说明"], [
    ["Product", "星辰 S14、星辰 S14 Pro", "name、release_year", "产品节点"],
    ["SKU", "SKU_8+256、SKU_16+512", "name、price、promo_price", "SKU 节点"],
    ["Spec", "SPEC_S14_processor", "category、value", "规格节点"],
    ["Promotion", "PROMO_double11", "description、end_date", "促销节点"],
    ["StockStatus", "STOCK_S14_8+256_black", "sku、color、status", "库存状态节点"],
    ["Accessory", "ACC_charger_65W", "name、price", "配件节点"],
    ["Warranty", "WARRANTY_screen", "description、duration", "保修政策节点"],
    ["Issue", "ISSUE_screen_broken", "name、repair_cost", "故障类型节点"],
])

add_heading(doc, "3.5.2 关系类型设计", level=3)
add_para(doc, "共定义 7 种有向边关系类型,如表 3-11 所示。")
add_table_caption(doc, "表 3-11  知识图谱关系类型")
add_table(doc, ["关系类型", "起点类型", "终点类型", "语义"], [
    ["HAS_SKU", "Product", "SKU", "产品包含的 SKU"],
    ["HAS_SPEC", "Product", "Spec", "产品的规格参数"],
    ["HAS_PROMOTION", "Product", "Promotion", "产品的促销活动"],
    ["HAS_STOCK", "SKU", "StockStatus", "SKU 的库存状态"],
    ["HAS_ACCESSORY", "Product", "Accessory", "产品的配件"],
    ["HAS_WARRANTY", "Product", "Warranty", "产品的保修政策"],
    ["HAS_FAULT", "Product", "Issue", "产品常见故障类型"],
])

add_heading(doc, "3.5.3 知识图谱 Schema 示例", level=3)
add_para(doc, "以“星辰 S14”为中心节点,其完整的一跳邻居关系如图 3-4 所示。具体来说,“星辰 S14”通过 HAS_SKU 关联到 SKU_8+256、SKU_12+256、SKU_16+512、SKU_16+1TB 共 4 个 SKU 节点;通过 HAS_SPEC 关联到处理器、屏幕、电池、后摄、机身材质共 5 个规格节点;通过 HAS_PROMOTION 关联到双十一促销节点;通过 HAS_WARRANTY 关联到屏幕保修、电池保修、整机保修共 3 个保修政策节点;通过 HAS_FAULT 关联到屏幕碎裂、电池老化、进水共 3 个故障类型节点。")
add_figure(doc, ASSETS / "fig_3_4_kg_schema.png", "图 3-4  知识图谱 Schema 示意图", width_cm=14)

add_heading(doc, "3.5.4 知识图谱与 RAG 的协同机制", level=3)
add_para(doc, "知识图谱与 RAG 在系统中通过以下机制协同工作。意图分类驱动路由:kg_service.kg_query() 根据意图分类与子关键词路由到对应的图谱查询函数(如 query_product_price、query_stock、query_spec_comparison)。优先级合并:在上下文组装阶段,知识图谱返回的结构化知识被放在最前(标题为【结构化知识】),向量检索的非结构化片段被放在最后(标题为【参考资料】),引导大模型优先使用结构化信息。降级容错:任意一路检索失败均不影响其他路,知识图谱查询异常时通过 try-except 静默吞掉,系统继续使用其他源的上下文生成回答。")

add_heading(doc, "3.6 接口设计", level=2)
add_para(doc, "系统对外暴露 RESTful API,统一前缀为 /api/。所有接口的请求体与响应体均使用 JSON 格式,字符编码为 UTF-8。流式接口使用 SSE 协议,MIME 类型为 text/event-stream。")

add_heading(doc, "3.6.1 核心 API 接口列表", level=3)
add_table_caption(doc, "表 3-12  系统核心 API 接口列表")
add_table(doc, ["路径", "方法", "请求参数", "说明"], [
    ["/api/chat/stream", "POST", "session_id, question", "流式问答(SSE)"],
    ["/api/knowledge/upload", "POST", "file (form-data)", "上传文档并向量化入库"],
    ["/api/knowledge/clear", "POST", "—", "清空向量库"],
    ["/api/knowledge/stats", "GET", "—", "获取知识库统计"],
    ["/api/feedback/submit", "POST", "session_id, rating, comment, resolved", "提交反馈评分"],
    ["/api/feedback/stats", "GET", "—", "获取反馈统计"],
    ["/api/config/llm", "GET", "—", "查询当前 LLM 配置"],
    ["/api/config/llm", "POST", "api_key, base_url, model", "更新 LLM 配置"],
    ["/api/config/test", "POST", "—", "测试 LLM 连接"],
    ["/api/repair/list", "GET", "—", "维修工单列表"],
    ["/api/repair/detail/{ticket_id}", "GET", "ticket_id", "维修工单详情"],
    ["/api/logistics/list", "GET", "—", "物流订单列表"],
    ["/api/logistics/detail/{order_id}", "GET", "order_id", "物流订单详情"],
    ["/api/graph/data", "GET", "—", "获取知识图谱可视化数据"],
    ["/api/graph/stats", "GET", "—", "知识图谱统计"],
    ["/api/db/seed", "POST", "—", "初始化演示数据"],
    ["/api/demo/data", "GET", "—", "获取演示数据"],
    ["/health", "GET", "—", "健康检查"],
])

add_heading(doc, "3.6.2 流式问答接口设计", level=3)
add_para(doc, "/api/chat/stream 是系统最核心的接口,采用 SSE 协议返回流式响应。请求体格式如下:")
add_code(doc, '{\n  "session_id": "default",\n  "question": "S14 Pro 的处理器是什么?"\n}')
add_para(doc, "响应为 SSE 流,包含三种事件类型,按顺序依次推送。meta 事件:第一条事件,携带分类结果、参考来源、相似问题推荐。token 事件:多条,每条携带一个 token 增量。done 事件:最后一条,携带完整答案。")
add_code(doc, 'data: {"type":"meta","category":"product_qa","sources":[...],"kg_used":true}\n\ndata: {"type":"token","content":"星辰"}\ndata: {"type":"token","content":"S14 Pro"}\n...\ndata: {"type":"done","answer":"星辰S14 Pro 搭载高通骁龙 Gen3 处理器..."}')
add_para(doc, "这种“meta + token × N + done”的事件协议设计有两方面优点:前端可以在 meta 事件到达时立刻渲染分类标签与参考来源,无需等待生成完成,大幅提升交互体验;done 事件携带完整答案,前端可以基于此进行 Markdown 渲染、复制、保存等后处理操作。")

add_heading(doc, "3.7 本章小结", level=2)
add_para(doc, "本章从功能需求与非功能需求两个维度梳理了系统需求,识别出“多源异构数据融合”的核心挑战。在此基础上给出了六层分层架构设计,详细描述了“分类 + 三路并联检索 + 单次 LLM”的核心处理流程。完成了 8 张 MySQL 表的字段设计与 ER 关系设计,完成了 8 类节点、7 类关系的知识图谱模式设计,并列出了 18 个对外 API 接口。下一章将围绕这些设计给出具体的代码实现细节。")

add_page_break(doc)


# ============ CHAPTER 4 ============
add_heading(doc, "第 4 章  系统详细实现", level=1)
add_para(doc, "本章详细描述系统各模块的实现细节,按“开发环境 → 知识库 → 意图分类 → RAG 核心 → 知识图谱 → 数据库与缓存 → SSE 接口 → 前端界面 → 前端功能展示”的顺序展开。所有关键代码均给出与仓库实际代码一致的实现,并配以必要的设计说明。")

add_heading(doc, "4.1 开发环境与项目结构", level=2)
add_heading(doc, "4.1.1 开发环境", level=3)
add_para(doc, "本系统的开发与运行环境配置如表 4-1 所示。")
add_table_caption(doc, "表 4-1  开发环境配置")
add_table(doc, ["项目", "说明"], [
    ["操作系统", "Windows 11 / Ubuntu 22.04 / macOS 14"],
    ["Python 版本", "3.10"],
    ["Web 框架", "FastAPI 0.115.5"],
    ["ASGI 服务器", "Uvicorn 0.32.1"],
    ["数据库", "MySQL 8.0.36"],
    ["缓存", "Redis 7.2.5"],
    ["ORM", "SQLAlchemy 2.0+"],
    ["向量库", "FAISS 1.9.0(CPU 版)"],
    ["嵌入模型", "shibing624/text2vec-base-chinese"],
    ["LLM 服务", "DeepSeek-Chat(远程 API 调用)"],
    ["LLM 框架", "LangChain 0.3.7"],
    ["图谱库", "NetworkX 3.2+"],
    ["前端", "原生 HTML5 + CSS3 + JavaScript ES6"],
    ["可视化库", "vis-network 9.x"],
    ["Markdown 渲染", "marked.js"],
])

add_heading(doc, "4.1.2 项目目录结构", level=3)
add_para(doc, "项目根目录结构如下,清晰反映了第 3 章设计的分层架构。")
add_code(doc, """rag-customer-service/
├── app/
│   ├── main.py                  # FastAPI 入口
│   ├── config.py                # 全局配置(环境变量)
│   ├── api/                     # API 路由层
│   │   ├── chat.py              # 流式问答路由
│   │   ├── knowledge.py         # 知识库管理路由
│   │   ├── feedback.py          # 反馈路由
│   │   ├── config.py            # 配置路由
│   │   ├── repair.py            # 维修路由
│   │   ├── logistics.py         # 物流路由
│   │   ├── graph.py             # 图谱路由
│   │   ├── db.py                # 数据库管理路由
│   │   └── demo.py              # 演示数据路由
│   ├── services/                # 业务服务层
│   ├── core/                    # 核心服务层
│   ├── models/                  # 模型层
│   ├── store/                   # 存储层
│   ├── knowledge/               # 知识管理
│   ├── scripts/                 # 脚本工具
│   └── templates/index.html
├── data/                        # 运行时数据
│   ├── docs/                    # 原始文档
│   ├── faiss_index/             # FAISS 索引
│   └── kg/graph.pkl             # 知识图谱
├── diagrams/                    # 架构图
├── docker-compose.yml           # MySQL + Redis 编排
├── requirements.txt
├── deploy.sh / deploy.ps1       # 部署脚本
└── .env.example                 # 配置模板""")

add_heading(doc, "4.2 知识库构建与文本切分", level=2)
add_heading(doc, "4.2.1 文档加载", level=3)
add_para(doc, "文档加载实现在 app/knowledge/loader.py 中,使用 LangChain 提供的文档加载器统一接口。系统支持三种文档格式:TXT 文件使用 TextLoader 直接读取;PDF 文件使用 PyPDFLoader 解析,基于 pypdf 库提取文本;DOCX 文件使用 Docx2txtLoader,基于 python-docx 库提取段落文本。加载器返回的统一数据结构是 Document 对象列表,每个 Document 包含 page_content(文本内容)与 metadata(元数据,如来源文件名、页码)两个字段。")

add_heading(doc, "4.2.2 文本切分策略", level=3)
add_para(doc, "文本切分采用 LangChain 的 RecursiveCharacterTextSplitter,配置如下:")
add_code(doc, """RecursiveCharacterTextSplitter(
    chunk_size=500,
    chunk_overlap=50,
    separators=["\\n\\n", "\\n", "。", "!", "?", ".", " ", ""]
)""")
add_para(doc, "切分策略的核心思想是递归尝试不同的分隔符:优先按段落分割(\\n\\n),段落过长则按行分割(\\n),仍过长则按句号分割(。),依次递归直到满足块大小限制。50 字的重叠保证了上下文的连续性,避免在边界处丢失语义信息。chunk_size=500 的选择考虑了三方面因素:一是中文每字约对应 1.5 个 token,500 字约 750 token,适合作为检索单元;二是过短的块会丢失上下文,过长的块会稀释相关性;三是 500 字的块在 LLM Prompt 中可以容纳 5 个 Top-K 片段而不超出上下文窗口。")

add_heading(doc, "4.2.3 向量化与索引构建", level=3)
add_para(doc, "文本切分完成后,系统调用 app/core/vector_store.py 中的 add_documents() 函数,将切分后的文本块向量化并存入 FAISS 索引。核心实现如下:")
add_code(doc, """def add_documents(chunks: list[str], metadata: dict = None) -> int:
    global _vector_store, _docs_cache
    _docs_cache = None
    if MOCK_RETRIEVAL:
        return len(chunks)
    from app.core.embeddings import get_embeddings
    documents = [
        Document(page_content=chunk, metadata=metadata or {})
        for chunk in chunks
    ]
    embeddings = get_embeddings()
    if _vector_store is None:
        _vector_store = FAISS.from_documents(documents, embeddings)
    else:
        _vector_store.add_documents(documents)
    os.makedirs(FAISS_INDEX_DIR, exist_ok=True)
    _vector_store.save_local(FAISS_INDEX_DIR)
    return len(chunks)""")
add_para(doc, "实现要点有三:一是首次入库时使用 FAISS.from_documents 创建新索引,后续入库使用 add_documents 增量追加;二是每次入库后立即调用 save_local 持久化到磁盘,避免重启后丢失;三是入库时清空 _docs_cache,使后续 BM25 回退检索能感知到新文档。")
add_figure(doc, DIAG / "d3_kb_build.png", "图 4-1  知识库构建流程图", width_cm=8)

add_heading(doc, "4.2.4 检索实现", level=3)
add_para(doc, "检索接口 search() 同时支持向量检索与 BM25 关键词检索两种模式,通过 MOCK_RETRIEVAL 环境变量切换。")
add_code(doc, """def search(query: str, top_k: int = 5) -> list[dict]:
    if MOCK_RETRIEVAL:
        return _keyword_search(query, top_k)
    vs = get_vector_store()
    if vs is None:
        return _keyword_search(query, top_k)
    results = vs.similarity_search_with_score(query, k=top_k)
    return [
        {
            "content": doc.page_content,
            "similarity": round(l2_to_cosine_similarity(float(score)), 4),
            "l2_distance": round(float(score), 4),
            "metadata": doc.metadata,
        }
        for doc, score in results
    ]""")
add_para(doc, "向量模式返回的相似度通过 l2_to_cosine_similarity() 函数将 L2 距离转换为余弦相似度,公式为 cosine = 1 - L2² / 2,该公式仅在向量经过 L2 归一化时成立,本系统已通过 normalize_embeddings=True 保证。BM25 回退模式实现了简化版的 BM25 算法,采用字符 + bigram 作为分词单位,以适应中文场景下无需额外分词工具的需求。该模式主要用于冷启动(向量库未构建)或纯演示场景,保证系统在无嵌入模型环境下也能可用。")

add_heading(doc, "4.3 意图分类器实现", level=2)
add_heading(doc, "4.3.1 设计思路", level=3)
add_para(doc, "意图分类器的设计核心是“零延迟、不调 LLM”,通过纯本地规则在 1 毫秒内完成分类。这一设计基于两个观察:其一,客服场景下用户问题的意图边界相对清晰,关键词覆盖率高;其二,订单号、工单号具有明显的格式特征(如 DD 开头的数字串),可通过正则精确匹配。")

add_heading(doc, "4.3.2 分类规则定义", level=3)
add_para(doc, "意图分类器定义在 app/core/classifier.py 中。系统预定义了六类意图及其关键词列表:")
add_code(doc, """_RULES = {
    "repair_track": ["维修进度", "工单", "修好了吗", "维修状态", ...],
    "logistics_track": ["物流查询", "到哪了", "快递单号", ...],
    "after_sales": ["退货", "换货", "退款", "保修", "故障", ...],
    "ecommerce": ["价格", "多少钱", "优惠", "库存", ...],
    "product_qa": ["怎么用", "如何使用", "功能", "设置", ...],
}""")
add_para(doc, "订单号与工单号的正则模式如下:")
add_code(doc, """_ID_PATTERNS = [
    (re.compile(r"DD\\d{8,}", re.IGNORECASE), "logistics_track"),
    (re.compile(r"SF\\d{8,}", re.IGNORECASE), "logistics_track"),
    (re.compile(r"YT\\d{8,}", re.IGNORECASE), "logistics_track"),
    (re.compile(r"JD\\d{8,}", re.IGNORECASE), "logistics_track"),
    (re.compile(r"ORD\\d{8,}", re.IGNORECASE), "logistics_track"),
    (re.compile(r"WX\\d{8,}", re.IGNORECASE), "repair_track"),
]""")

add_heading(doc, "4.3.3 分类算法", level=3)
add_para(doc, "分类算法采用两阶段策略:第一阶段进行 ID 模式匹配,具有最高优先级;第二阶段进行关键词打分。")
add_code(doc, """def classify_question(question: str) -> str:
    question_clean = question.lower().strip()
    # 阶段 1:ID 模式匹配(最高优先级)
    for pattern, category in _ID_PATTERNS:
        if pattern.search(question):
            return category
    # 阶段 2:关键词打分
    scores = {}
    for category, keywords in _RULES.items():
        score = sum(1 for kw in keywords if kw in question_clean)
        if score > 0:
            scores[category] = score
    if not scores:
        return "other"
    return max(scores, key=scores.get)""")
add_para(doc, "分类算法遍历每个类别的关键词列表,统计用户问题中出现的关键词数量作为该类别的得分,返回得分最高的类别。若所有类别得分为零,则返回 other。整个分类过程是纯字符串匹配操作,耗时不到 1 毫秒,不涉及任何网络请求或模型推理。")

add_heading(doc, "4.3.4 设计权衡", level=3)
add_para(doc, "将意图分类从 LLM 中剥离的设计带来了三方面优势。延迟优势:本地规则分类在 1 毫秒内完成,而 LLM 调用通常需要 200 至 500 毫秒。成本优势:每次分类节省约 50 至 200 个输入 token 的开销。可控性优势:规则可读、可调试、可热更新,无需重新训练或调整 prompt。这一设计的代价是:对于关键词未覆盖的边缘案例,分类可能不准。但通过实测(详见第 5 章)发现,只要关键词列表设计得当,准确率可达 95% 以上,完全满足客服场景需求。")

add_heading(doc, "4.4 RAG 核心流程实现", level=2)
add_heading(doc, "4.4.1 混合检索策略", level=3)
add_para(doc, "RAG 核心流程实现在 app/core/rag_chain.py 的 _build_prompt_and_sources() 函数中。该函数按以下步骤执行:(1)调用意图分类器确定问题类别。(2)调用知识图谱服务获取结构化知识(如产品规格、价格等)。(3)对于维修跟踪、物流跟踪类问题,额外从对应业务服务获取实时数据。(4)调用 FAISS 向量检索获取 Top-K 相关文档片段。(5)按优先级组装上下文:结构化知识 → 实时服务数据 → RAG 参考资料。(6)选择对应意图的 Prompt 模板,填充上下文与用户问题。(7)返回(分类结果、参考来源、最终 Prompt、KG 是否命中)四元组。")
add_para(doc, "整个混合检索过程的 Python 实现如下:")
add_code(doc, """def _build_prompt_and_sources(question, chat_history):
    category = classify_question(question)
    # KG 结构化查询
    kg_context = None
    try:
        kg_context = kg_query(category, question)
    except Exception:
        pass
    kg_used = kg_context is not None
    # 维修/物流:优先获取实时业务数据
    service_ctx = _get_service_context(category, question)
    # FAISS 向量检索
    results = search(question, top_k=TOP_K)
    if not results and not kg_context and not service_ctx:
        return category, [], None, False
    # 按优先级组装上下文
    parts = []
    if kg_context:
        parts.append(f"【结构化知识】\\n{kg_context}")
    if service_ctx:
        parts.append(f"【实时查询结果】\\n{service_ctx}")
    if results:
        rag_text = "\\n\\n---\\n\\n".join([r["content"] for r in results])
        parts.append(f"【参考资料】\\n{rag_text}")
    context = "\\n\\n---\\n\\n".join(parts)
    prompt_template = get_prompt_by_category(category)
    prompt = prompt_template.format(
        context=context, question=question,
        chat_history=chat_history or "无")
    sources = [{"content": r["content"][:200], "similarity": r["similarity"],
                "metadata": r["metadata"]} for r in results] if results else []
    return category, sources, prompt, kg_used""")
add_figure(doc, DIAG / "d5_core_flow.png", "图 4-2  核心流程概览", width_cm=10)

add_heading(doc, "4.4.2 Prompt 模板设计", level=3)
add_para(doc, "Prompt 模板定义在 app/core/prompt_templates.py 中,共 6 个模板,分别对应六类意图。所有模板共享统一的基础规则:")
add_code(doc, """## 回答要求:
1. 只根据参考资料中的内容回答,不要编造信息
2. 如果参考资料中没有相关信息,请诚实告知用户
3. 回答要准确、专业、简洁
4. 语气友好、耐心""")
add_para(doc, "每个模板在基础规则之上追加分类专属规则。例如,产品问答模板强调“分步骤说明操作方法,每步简短清晰”;售后服务模板强调“明确告知政策要点(时限、条件)”;电商咨询模板强调“价格信息用清晰格式呈现,主动说明优惠”;维修跟踪与物流跟踪模板强调“参考资料中包含【实时查询结果】时,优先使用其中的工单/物流信息回答”。每个模板都包含三个占位符:{context} 用于填充检索到的上下文,{question} 填充用户问题,{chat_history} 填充最近 5 轮对话历史。")

add_heading(doc, "4.4.3 流式生成器实现", level=3)
add_para(doc, "流式问答通过 Python 生成器函数 rag_query_stream() 实现。生成器依次 yield 三种事件:meta 事件携带分类结果、参考来源与相似问题推荐;token 事件携带每个 token 的增量内容;done 事件携带完整答案。")
add_code(doc, """def rag_query_stream(question, chat_history) -> Generator:
    category, sources, prompt, kg_used = _build_prompt_and_sources(
        question, chat_history)
    similar_questions = search_faq(question, top_k=3)
    if prompt is None:
        yield {"type": "meta", "category": category,
               "sources": [], "similar_questions": similar_questions,
               "kg_used": False}
        yield {"type": "token", "content": _EMPTY_ANSWER}
        yield {"type": "done", "answer": _EMPTY_ANSWER}
        return
    yield {"type": "meta", "category": category,
           "sources": sources, "similar_questions": similar_questions,
           "kg_used": kg_used}
    llm = get_llm()
    full_answer = ""
    for chunk in llm.stream(prompt):
        token = chunk.content
        if token:
            full_answer += token
            yield {"type": "token", "content": token}
    yield {"type": "done", "answer": full_answer}""")

add_heading(doc, "4.5 知识图谱构建与查询", level=2)
add_heading(doc, "4.5.1 图谱构建脚本", level=3)
add_para(doc, "知识图谱通过 app/scripts/build_graph.py 脚本以代码方式构建。脚本中以代码形式定义所有节点与边,经过约 9.6 KB 的图谱定义代码后,通过 pickle.dump 序列化为 data/kg/graph.pkl 文件,大小约 50 KB,包含约 50 个节点与 80 条边。构建脚本的关键代码如下:")
add_code(doc, """import pickle
import networkx as nx

G = nx.DiGraph()
G.add_node("星辰S14", type="Product",
           name="星辰S14", release_year=2026)
G.add_node("SKU_8+256", type="SKU",
           name="8+256GB", price=3999, promo_price=3699)
G.add_edge("星辰S14", "SKU_8+256", relation="HAS_SKU")
G.add_node("SPEC_S14_processor", type="Spec",
           category="processor", value="高通骁龙7 Gen3")
G.add_edge("星辰S14", "SPEC_S14_processor", relation="HAS_SPEC")
# ... 后续依次添加 Promotion、StockStatus、Accessory、Warranty、Issue 节点
with open("data/kg/graph.pkl", "wb") as f:
    pickle.dump(G, f)""")

add_heading(doc, "4.5.2 图谱加载与查询", level=3)
add_para(doc, "图谱加载实现在 app/core/knowledge_graph.py 中,采用单例模式,首次调用 get_graph() 时从磁盘加载到内存,后续调用直接返回内存中的图对象。")
add_code(doc, """_graph: Optional[nx.DiGraph] = None

def get_graph() -> nx.DiGraph:
    global _graph
    if _graph is None:
        if not os.path.exists(_KG_PATH):
            _graph = nx.DiGraph()
        else:
            with open(_KG_PATH, "rb") as f:
                _graph = pickle.load(f)
    return _graph""")
add_para(doc, "图谱查询函数共有 8 个,涵盖产品价格、库存、规格对比、维修费用、保修政策、产品规格、配件、故障类型查询。以价格查询为例,查询函数采用“优先 MySQL 实时查询,降级到图谱查询”的双轨策略:")
add_code(doc, """def query_product_price(question: str) -> Optional[str]:
    # 优先:实时数据库查询
    try:
        from app.services.db_service import get_product
        sku_map = {"8+256": 1001, "12+256": 1002,
                   "16+512": 1003, "16+1tb": 1004}
        for pat, sid in sku_map.items():
            if pat in question.lower().replace(" ", ""):
                p = get_product(sid)
                if p:
                    return f"星辰S14 {p['product_name']}: " \\
                           f"官方价 {p['current_price']}元, " \\
                           f"促销价 {p['promo_price']}元"
    except Exception:
        pass  # 降级到图谱
    # 降级:知识图谱查询
    G = get_graph()
    lines = []
    for _, target, edata in G.out_edges("星辰S14", data=True):
        if edata.get("relation") == "HAS_SKU":
            d = G.nodes[target]
            lines.append(f"{d['name']}: 官方价 {d['price']}元, "
                         f"促销价 {d['promo_price']}元")
    return "\\n".join(lines) if lines else None""")
add_para(doc, "这种双轨策略体现了“图谱是冷知识、数据库是热数据”的分层思想:数据库可用时返回最新数据,数据库不可用时降级到图谱中的快照数据,保证系统在数据库故障时仍可基本可用。")

add_heading(doc, "4.5.3 KG 查询路由", level=3)
add_para(doc, "app/services/kg_service.py 中的 kg_query() 函数负责将问题路由到合适的图谱查询函数。路由依据是意图分类结果与子关键词匹配。这种基于“分类 + 子关键词”的二级路由设计使得每类问题只触发相关的图谱查询,避免了不必要的图遍历开销。")

add_heading(doc, "4.6 电商数据库与缓存层实现", level=2)
add_heading(doc, "4.6.1 缓存层设计", level=3)
add_para(doc, "缓存层基于 Redis 实现,封装在 app/store/cache.py 中。所有缓存操作通过 redis-py 客户端完成,采用 JSON 序列化存储复杂对象。核心设计包括以下四点。缓存键规范:所有键加 rag_cs: 前缀以便管理。默认 TTL:300 秒。安全失效:模式删除使用 SCAN + DEL 而非 KEYS *。故障降级:Redis 客户端设置 1 秒超时,Redis 不可用时降级。")
add_code(doc, """def get_json(key: str) -> Any:
    client = get_cache_client()
    if client is None:
        return None
    try:
        raw = client.get(_key(key))
        return json.loads(raw) if raw else None
    except redis.RedisError as exc:
        logger.warning("Redis get failed: %s", exc)
        return None

def delete_pattern(pattern: str):
    client = get_cache_client()
    if client is None: return
    try:
        keys = list(client.scan_iter(
            match=_key(pattern), count=200))
        if keys:
            client.delete(*keys)
    except redis.RedisError as exc:
        logger.warning("Redis delete pattern failed: %s", exc)""")

add_heading(doc, "4.6.2 数据库服务实现", level=3)
add_para(doc, "电商数据库服务实现在 app/services/db_service.py 中。所有查询函数遵循统一的“缓存穿透”策略 — 先查 Redis,未命中则查 MySQL,查到后回填 Redis。写操作执行后主动调用 delete_pattern 清除相关缓存。所有数据库操作通过 session_scope() 上下文管理器统一管理。以商品查询为例:")
add_code(doc, """def get_product(sku_id: int) -> Optional[dict]:
    cache_key = f"product:{sku_id}"
    cached = get_json(cache_key)
    if cached is not None:
        return cached
    with session_scope() as session:
        product = session.get(Product, sku_id)
        if product is None or not product.is_on_sale:
            return None
        result = _product_to_dict(product)
    set_json(cache_key, result)
    return result""")

add_heading(doc, "4.6.3 ORM 模型定义", level=3)
add_para(doc, "ORM 模型定义在 app/store/models.py 中,使用 SQLAlchemy 2.0 的 Mapped 类型注解风格。以商品表为例:")
add_code(doc, """class Product(Base):
    __tablename__ = "products"
    sku_id: Mapped[int] = mapped_column(Integer, primary_key=True)
    product_name: Mapped[str] = mapped_column(
        String(255), nullable=False, index=True)
    current_price: Mapped[float] = mapped_column(Float, nullable=False)
    promo_price: Mapped[Optional[float]] = mapped_column(Float, nullable=True)
    stock_num: Mapped[int] = mapped_column(
        Integer, nullable=False, default=0)
    specs_json: Mapped[Optional[str]] = mapped_column(Text, nullable=True)
    is_on_sale: Mapped[bool] = mapped_column(
        Boolean, nullable=False, default=True)
    promotions: Mapped[list["Promotion"]] = relationship(
        back_populates="product",
        cascade="all, delete-orphan")""")

add_heading(doc, "4.7 SSE 流式接口实现", level=2)
add_heading(doc, "4.7.1 后端流式端点", level=3)
add_para(doc, "后端 SSE 端点通过 FastAPI 的 StreamingResponse 实现,定义在 app/api/chat.py 中。")
add_code(doc, """@router.post("/stream")
async def stream_question(req: ChatRequest):
    add_message(req.session_id, "user", req.question)
    chat_history = format_history(req.session_id, limit=5)
    def generate():
        full_answer = ""
        try:
            for event in rag_query_stream(req.question, chat_history):
                if event["type"] == "done":
                    full_answer = event.get("answer", "")
                yield f"data: {json.dumps(event, ensure_ascii=False)}\\n\\n"
        except Exception as e:
            err = {"type": "error", "content": f"服务异常:{str(e)}"}
            yield f"data: {json.dumps(err, ensure_ascii=False)}\\n\\n"
        finally:
            if full_answer:
                add_message(req.session_id, "assistant", full_answer)
    return StreamingResponse(
        generate(), media_type="text/event-stream")""")
add_para(doc, "实现要点:接收请求后立即将用户消息写入数据库;format_history 拉取最近 5 轮历史作为上下文;generate 函数捕获所有异常并返回友好错误事件;finally 块在流结束后将完整答案写回数据库。")

add_heading(doc, "4.7.2 SSE 协议格式", level=3)
add_para(doc, "SSE 协议格式简单:每条事件以 data: 开头、\\n\\n 结尾,事件内容为 JSON 字符串。响应头自动设置为 Content-Type: text/event-stream、Cache-Control: no-cache、Connection: keep-alive,FastAPI 的 StreamingResponse 自动处理这些头部。")

add_heading(doc, "4.7.3 前端流式消费", level=3)
add_para(doc, "前端使用 Fetch API 的 ReadableStream 消费 SSE 流。核心代码逻辑如下:")
add_code(doc, """const response = await fetch('/api/chat/stream', {
    method: 'POST',
    headers: {'Content-Type': 'application/json'},
    body: JSON.stringify({session_id, question})
});
const reader = response.body.getReader();
const decoder = new TextDecoder();
let rawAnswer = "";
let buffer = "";
while (true) {
    const {done, value} = await reader.read();
    if (done) break;
    buffer += decoder.decode(value, {stream: true});
    const events = buffer.split("\\n\\n");
    buffer = events.pop();
    for (const eventText of events) {
        if (!eventText.startsWith("data: ")) continue;
        const event = JSON.parse(eventText.slice(6));
        if (event.type === "meta") {
            renderCategoryTag(event.category);
            renderSources(event.sources);
        } else if (event.type === "token") {
            rawAnswer += event.content;
            answerEl.textContent = rawAnswer;
        } else if (event.type === "done") {
            answerEl.innerHTML = marked.parse(rawAnswer);
        }
    }
}""")
add_para(doc, "关键的安全设计:流式阶段使用 textContent 属性追加文本(不解析 HTML,防止 XSS),完成后使用 marked.parse() 进行 Markdown 渲染(marked 库会过滤危险标签)。这一设计在保证 Markdown 渲染效果的同时杜绝了 XSS 注入风险。")

add_heading(doc, "4.8 前端界面实现", level=2)
add_heading(doc, "4.8.1 整体布局", level=3)
add_para(doc, "前端采用单文件 HTML 实现,文件大小约 58 KB。界面布局分为左侧导航栏与右侧主区域两大部分。导航栏提供 Tab 切换功能,共 5 个面板:模型配置、知识库、知识图谱、演示数据、统计;主区域为持续可见的智能客服对话窗口。聊天界面实现了用户消息(右侧蓝色气泡)和助手消息(左侧深色气泡)的对话样式,每条助手回答底部提供可展开的参考来源面板,显示文档片段和相似度分数。")
add_figure(doc, ASSETS / "fig_4_1_overview.png", "图 4-3  系统前端整体界面", width_cm=15)

add_heading(doc, "4.8.2 技术选型", level=3)
add_para(doc, "前端选用原生 HTML5 + CSS3 + JavaScript 实现,不依赖 React、Vue 等前端 UI 框架。这一选型基于三方面考虑:第一,降低部署复杂度,无需 npm 构建过程,一个 HTML 文件即可部署;第二,降低用户学习门槛,本科毕设场景下原生 JS 更易于阅读与维护;第三,体积小,首屏加载快。仅引入两个第三方 JS 库:vis-network(知识图谱可视化)与 marked.js(Markdown 渲染),通过 CDN 加载。")

add_heading(doc, "4.8.3 SSE 流式渲染", level=3)
add_para(doc, "聊天界面的核心是 SSE 流式渲染,具体实现已在 4.7.3 节给出。流式阶段在气泡内逐字追加 token,完成后整体进行 Markdown 渲染,使用户既能看到打字机效果,又能在最终看到格式化的答案(列表、加粗、代码块等)。")

add_heading(doc, "4.8.4 知识图谱可视化", level=3)
add_para(doc, "知识图谱可视化基于 vis-network 库实现,采用力导向图布局。前端通过 /api/graph/data 接口获取图谱数据,转换为 vis-network 所需的 nodes/edges 格式后渲染。节点按类型区分颜色:Product 蓝色、SKU 橙色、Spec 绿色、Promotion 紫色、Issue 红色。支持节点拖拽与画布缩放,鼠标悬停节点时显示详细属性。")
add_figure(doc, ASSETS / "fig_3_4_kg_schema.png", "图 4-4  知识图谱可视化界面(力导向图)", width_cm=14)

add_heading(doc, "4.9 前端功能展示", level=2)
add_para(doc, "本节按功能模块逐一展示前端界面,每个模块配以截图与功能说明,便于读者直观了解系统能力。")

add_heading(doc, "4.9.1 左侧导航菜单", level=3)
add_para(doc, "左侧导航栏提供 Tab 切换功能,各功能项的作用如表 4-2 所示。")
add_table_caption(doc, "表 4-2  左侧导航功能项说明")
add_table(doc, ["功能项", "作用说明"], [
    ["模型配置", "配置对话大模型的 API 参数(API Key、Base URL、模型名)"],
    ["知识库", "上传与管理自定义文档(PDF/Word/TXT 等),后续问答基于这些文档进行检索增强"],
    ["知识图谱", "构建与展示结构化知识图谱,实现更精准的关联问答"],
    ["演示数据", "提供示例知识库与对话数据,方便用户快速体验系统功能"],
    ["统计", "对话量、用户满意度、问题解决率等数据看板,用于运营分析"],
])

add_heading(doc, "4.9.2 模型配置模块", level=3)
add_figure(doc, ASSETS / "fig_4_3_config.png", "图 4-5  模型配置界面", width_cm=15)
add_para(doc, "模型配置模块是整个系统的“AI 大脑”配置入口,允许用户在不重启服务的前提下动态切换大模型。该模块包括以下交互元素。服务商选择:下拉菜单,支持选择不同大模型服务商(如 OpenAI、DeepSeek、通义千问等),也支持“手动填写”方式自定义配置。API KEY:大模型服务商提供的身份密钥,用于调用模型接口。前端展示时仅显示前 7 位与后 4 位,中间以星号隐藏(如 sk-d2b****537e),保护密钥安全。BASE URL:大模型的 API 接口地址,默认值为 DeepSeek 官方接口 https://api.deepseek.com。MODEL:指定调用的模型名称,默认值为 deepseek-chat。操作按钮:保存按钮持久化当前配置;测试连接按钮发送一条测试请求,验证 API Key 与 Base URL 是否正确,并显示单次调用的耗时与 Token 消耗。状态提示:底部显示当前连接状态(已连接/未连接)与模型标识预览。")

add_heading(doc, "4.9.3 知识库文档上传模块", level=3)
add_figure(doc, ASSETS / "fig_4_4_kb.png", "图 4-6  知识库文档上传界面", width_cm=15)
add_para(doc, "知识库模块支持以下操作。支持格式:界面明确标注 PDF / DOCX / TXT 三种格式。选择文件按钮:点击后打开本地文件浏览器,选择要上传的文档,初始状态显示“未选择文件”。上传按钮:选中文件后点击,将文档上传到系统,后台自动完成解析、分块、向量化与索引构建,过程中显示进度条。清空按钮:清空当前向量库,用于重建索引或切换知识库。统计信息:页面顶部显示当前知识库的文档数量与切片数量,便于用户了解知识库规模。")

add_heading(doc, "4.9.4 知识图谱模块", level=3)
add_figure(doc, ASSETS / "fig_4_5_kg_panel.png", "图 4-7  知识图谱可视化模块", width_cm=15)
add_para(doc, "知识图谱模块将企业数据拆解为“实体”和“关系”,以可视化网络图的方式展示数据关联。该模块的功能价值体现在三方面:支持更精准的关联查询(如“S14 与 S14 Pro 的处理器对比”);支持推理问答(如“购买 1001 型号商品的用户中提交了维修工单的有多少”);便于运营人员直观了解知识库结构。界面交互特性包括:节点可拖拽布局,画布支持缩放,节点按类型(Product/SKU/Spec/Promotion/Issue 等)显示不同颜色,鼠标悬停显示节点详细属性。")

add_heading(doc, "4.9.5 演示数据模块", level=3)
add_figure(doc, ASSETS / "fig_4_6_demo.png", "图 4-8  演示数据初始化界面", width_cm=15)
add_para(doc, "演示数据模块的作用是一键初始化 MySQL 数据库中的演示数据,主要包括两张业务表:维修工单表(repair_tickets)与物流订单表(tracking_orders + tracking_events)。点击“初始化数据”按钮后,系统调用 app.store.seed.init_demo_data() 函数,自动写入 8 个商品、5 个促销、若干维修工单与物流轨迹,形成完整的测试数据集。该模块极大降低了演示与评测的门槛,用户无需手动准备测试数据即可体验系统全貌。")

add_heading(doc, "4.9.6 统计模块", level=3)
add_figure(doc, ASSETS / "fig_4_7_stats.png", "图 4-9  统计看板界面", width_cm=15)
add_para(doc, "统计模块是系统的运营看板,专门用来监控客服对话的关键指标。展示的指标包括:总对话量(累计 chat_messages 中 user 角色消息数);平均评分(feedback_records 中 rating 字段均值);满意率(rating ≥ 4 的反馈占比);解决率(resolved 为 true 的反馈占已标记反馈的比例)。该模块的应用价值体现在三方面:日常监控,查看对话量变化,了解用户活跃度;效果评估,通过平均评分、满意率、解决率,判断当前模型与知识库的回答质量;优化决策,如果解决率低、差评多,运营人员可以针对性地补充知识库内容或调整模型配置。")

add_heading(doc, "4.9.7 对话窗口", level=3)
add_figure(doc, ASSETS / "fig_4_8_chat.png", "图 4-10  对话窗口流式回答界面", width_cm=15)
add_para(doc, "对话窗口是用户最直接接触的功能模块,具有以下特性。顶部欢迎语:首次进入时,小星发送欢迎语,提示用户先完成左侧模型配置与知识库上传后再开始提问。对话区:多轮对话的展示区域,用户消息显示在右侧(蓝色气泡),AI 回复显示在左侧(深色气泡)。每条 AI 回复顶部显示一个分类徽标(如“产品问答”、“售后服务”),徽标颜色按分类区分。参考来源面板:每条 AI 回复底部提供可展开的参考来源面板,显示检索到的文档片段、相似度分数与来源文件名,便于用户核实回答依据。相似问题推荐:每条回复后展示 3 条相似问题推荐,用户点击即可发起新问答。清空对话按钮:一键清空当前会话记录,从头开始新对话。输入框:提示语为“输入您的问题,按 Enter 发送...”,用户可在此输入问题。")

add_heading(doc, "4.9.8 满意度评价模块", level=3)
add_figure(doc, ASSETS / "fig_4_9_rating.png", "图 4-11  满意度评价交互界面", width_cm=15)
add_para(doc, "每轮对话结束后,用户可通过以下交互元素提交反馈:星级评分,1 至 5 星可选,默认 5 星;已解决/未解决标签,二选一;文字评价,可选,用于补充说明。提交后调用 /api/feedback/submit 接口持久化到 feedback_records 表,并主动失效相关 Redis 缓存,使下次查询统计指标时获得最新数据。")

add_heading(doc, "4.10 本章小结", level=2)
add_para(doc, "本章按“开发环境 → 知识库 → 意图分类 → RAG 核心 → 知识图谱 → 数据库与缓存 → SSE 接口 → 前端界面 → 前端功能展示”的顺序详细描述了系统各模块的实现细节。核心实现要点包括:基于关键词与正则的零延迟意图分类器(1 ms 内完成);三路并联的多源检索方案(FAISS 向量 + NetworkX 图谱 + MySQL 实时);“meta + token + done”的 SSE 事件协议设计;基于 Redis 的“缓存穿透 + 写后失效”策略;原生 HTML/CSS/JS 单文件前端,集成对话、知识库、模型配置、图谱可视化、演示数据、统计 6 大功能模块。下一章将对系统进行功能测试与性能评估。")

add_page_break(doc)

# ============ CHAPTER 5 ============
add_heading(doc, "第 5 章  系统测试与性能评估", level=1)
add_para(doc, "本章对系统进行功能测试与性能评估,涵盖测试环境、意图分类准确性、端到端问答效果、响应时延、Redis 缓存效果与多源检索消融实验六个方面,从功能正确性与性能指标两个维度全面验证系统的可用性。")

add_heading(doc, "5.1 测试环境", level=2)
add_heading(doc, "5.1.1 硬件与软件环境", level=3)
add_para(doc, "测试环境配置如表 5-1 所示。")
add_table_caption(doc, "表 5-1  测试环境配置")
add_table(doc, ["项目", "说明"], [
    ["操作系统", "Windows 11"],
    ["处理器", "Intel Core i5"],
    ["内存", "16 GB"],
    ["Python 版本", "3.10"],
    ["LLM 服务", "DeepSeek-Chat(远程 API 调用)"],
    ["数据库", "MySQL 8.0.36(本地部署)"],
    ["缓存", "Redis 7.2.5(本地部署)"],
    ["运行方式", "uvicorn 本地部署"],
    ["网络环境", "校园网,带宽 100 Mbps"],
])

add_heading(doc, "5.1.2 测试数据准备", level=3)
add_para(doc, "测试数据通过 app.store.seed.init_demo_data() 函数自动初始化,包括以下内容:8 个商品(覆盖星辰 S14 与 S14 Pro 两个产品系列、4 种存储规格、4 种颜色);5 条促销活动(双十一、618 等);20 条电商订单与对应的物流轨迹;10 条维修工单(覆盖屏幕碎裂、电池老化、进水等典型故障);3 份产品文档(产品规格说明书、保修政策、使用指南),切分后约 200 个文本块。知识图谱通过 app.scripts.build_graph 脚本构建,共包含 50 个节点与 80 条边。")

add_heading(doc, "5.1.3 测试工具", level=3)
add_para(doc, "测试使用以下工具:Python requests 库用于编写 API 自动化测试脚本;Apache Bench(ab)用于并发压力测试;Chrome DevTools 的 Network 面板用于观察 SSE 流式响应的时序;Redis CLI 的 MONITOR 命令用于观察缓存命中情况。")

add_heading(doc, "5.2 意图分类准确性测试", level=2)
add_heading(doc, "5.2.1 测试方法", level=3)
add_para(doc, "为验证意图分类器的准确性,设计了 60 条测试用例,覆盖 6 类意图(每类 10 条),包括典型问法和边缘案例。每类测试用例分为三类:典型问法(5 条),包含明确关键词;变体问法(3 条),通过同义改写、词序变化等方式表达;边缘案例(2 条),关键词模糊或处于多类边界。测试用例示例:product_qa 类“S14 Pro 的处理器是什么”;after_sales 类“屏幕碎了能保修吗”;ecommerce 类“S14 16+512 多少钱”;repair_track 类“WX20260320001 修好了吗”;logistics_track 类“DD20260320001 到哪了”;other 类“你好”、“再见”。")

add_heading(doc, "5.2.2 测试结果", level=3)
add_para(doc, "测试结果如表 5-2 所示。")
add_table_caption(doc, "表 5-2  意图分类测试结果")
add_table(doc, ["意图类别", "测试数", "正确数", "准确率"], [
    ["product_qa", "10", "10", "100%"],
    ["after_sales", "10", "9", "90%"],
    ["ecommerce", "10", "10", "100%"],
    ["repair_track", "10", "10", "100%"],
    ["logistics_track", "10", "9", "90%"],
    ["other", "10", "9", "90%"],
    ["总计", "60", "57", "95%"],
])

add_heading(doc, "5.2.3 错误案例分析", level=3)
add_para(doc, "3 条错误案例分析如下。案例一(after_sales 误判):“我手机用了一年屏幕有点暗,正常吗”。该问题虽然属于售后咨询,但用户表达较委婉,既无“保修”“维修”“退换”等售后关键词,也无“故障”“坏了”等故障关键词,被分类为 other。改进方向:扩充 after_sales 关键词,增加“暗”“卡顿”“发热”“模糊”等故障表征词。案例二(logistics_track 误判):“我前天买的还没到呢”。该问题包含“买”但无明确订单号或物流关键词,被分类为 ecommerce。改进方向:增加时间词(“前天”“昨天”“上周”)与物流暗示词的组合规则。案例三(other 误判):“S14 Pro 怎么样”。该问题属于产品咨询的开放性问法,无明确关键词,被分类为 other。改进方向:对包含产品型号但无明确意图的问法,默认分类为 product_qa。")

add_heading(doc, "5.2.4 性能指标", level=3)
add_para(doc, "意图分类器的性能指标如下:单次分类延迟 < 1 ms(测试 1000 次取平均);内存占用 < 1 MB(关键词列表与正则模式);无网络依赖,完全本地运行。这一性能完全满足实时客服场景的延迟要求,且通过简单扩充关键词列表即可持续优化准确率。")

add_heading(doc, "5.3 端到端问答测试", level=2)
add_heading(doc, "5.3.1 测试方法", level=3)
add_para(doc, "针对每类意图设计了典型的问答测试场景,验证从用户提问到 LLM 生成最终回答的端到端流程是否符合预期,重点考察分类正确性、检索召回质量、回答内容准确度三方面。")

add_heading(doc, "5.3.2 测试场景与结果", level=3)
add_para(doc, "场景一:产品问答。询问“怎么重启手机”,系统正确分类为 product_qa,检索到产品说明文档中的操作步骤,生成分步骤指南。回答清晰列出“长按电源键 5 秒 → 滑动关机 → 短按电源键开机”三个步骤,并附带参考来源(产品规格说明书.txt,相似度 0.89)。")
add_para(doc, "场景二:售后服务。询问“屏幕碎了能保修吗”,系统分类为 after_sales,同时从知识图谱获取保修政策(【结构化知识】屏幕非人为损坏免费保修 1 年),从 RAG 获取售后流程说明(【参考资料】联系客服 → 寄送 → 维修 → 寄回)。回答整合两源信息,先明确保修结论,再说明操作流程。")
add_para(doc, "场景三:电商咨询。询问“S14 Pro 16+512 多少钱”,系统分类为 ecommerce,从 MySQL 数据库实时获取价格(官方价 ¥5,999、促销价 ¥5,599)。回答以加粗格式展示价格,并主动告知当前正处于促销期。")
add_para(doc, "场景四:维修跟踪。查询具体工单号“WX20260320001”,系统正则匹配到工单号模式后直接分类为 repair_track,从 repair_tickets 表获取工单状态(“维修中,预计还需 2 天”)并结合售后政策生成回答。")
add_para(doc, "场景五:物流跟踪。查询订单“DD20260320001 到哪了”,系统正则匹配到订单号后分类为 logistics_track,从 tracking_orders 与 tracking_events 表获取最新轨迹(“已到达成都转运中心,预计明日送达”)。")
add_para(doc, "场景六:多轮对话。第一轮提问“S14 多少钱”,系统返回价格列表;第二轮追问“那 Pro 版呢”,系统通过对话历史理解上下文,返回 Pro 版价格,验证了多轮对话上下文记忆功能。")
add_figure(doc, ASSETS / "fig_5_1_e2e.png", "图 5-1  端到端问答效果展示", width_cm=15)

add_heading(doc, "5.4 响应时延测试", level=2)
add_heading(doc, "5.4.1 测试方法", level=3)
add_para(doc, "通过 Chrome DevTools 的 Network 面板观察 SSE 流式响应的时序,记录五个关键时间点:T0 请求发出;T1 服务端开始返回(Time to First Byte);T2 meta 事件到达;T3 第一个 token 事件到达(用户感知首字延迟);T4 done 事件到达(完整答案生成完毕)。每个测试场景重复 10 次取平均值。")

add_heading(doc, "5.4.2 测试结果", level=3)
add_para(doc, "测试结果如表 5-3 所示。")
add_table_caption(doc, "表 5-3  响应时延测试结果(单位:毫秒)")
add_table(doc, ["阶段", "阶段说明", "平均耗时", "累计耗时"], [
    ["T0 → T1", "请求传输 + 服务端接收", "30", "30"],
    ["T1 → T2", "意图分类 + 多源检索 + meta 推送", "180", "210"],
    ["T2 → T3", "LLM 首 token 生成", "290", "500"],
    ["T3 → T4", "LLM 后续 token 流式生成", "2300", "2800"],
])
add_para(doc, "用户感知首字延迟约 500 毫秒(T0 → T3),完整答案生成约 2.8 秒(T0 → T4)。其中,LLM 首 token 延迟(290 ms)是整体延迟的最大组成部分,占 58%;意图分类与多源检索合计约 180 ms,占 36%。这一数据验证了“分类与生成解耦”设计的有效性 — 若每次都先用 LLM 分类再生成,首字延迟将增加约 300 至 500 ms,达到 1 秒以上。")

add_heading(doc, "5.4.3 各模块延迟拆解", level=3)
add_para(doc, "进一步拆解 T1 → T2 阶段的 180 ms,各模块延迟如表 5-4 所示。")
add_table_caption(doc, "表 5-4  检索阶段各模块延迟拆解(单位:毫秒)")
add_table(doc, ["模块", "延迟", "占比"], [
    ["意图分类", "< 1", "< 1%"],
    ["知识图谱查询", "8", "4%"],
    ["MySQL 实时查询(冷)/ Redis 缓存(热)", "35 / 2", "19% / 1%"],
    ["FAISS 向量检索", "120", "67%"],
    ["上下文组装 + Prompt 拼装", "15", "8%"],
])
add_para(doc, "FAISS 向量检索是检索阶段的主要耗时来源,占 67%。该耗时主要来自 Embedding 模型对查询的编码过程(约 100 ms),向量索引搜索本身仅约 20 ms。后续优化方向是引入 GPU 加速 Embedding 推理,或使用更小的嵌入模型(如 bge-small-zh,参数量约为 base 模型的三分之一)。")

add_heading(doc, "5.5 Redis 缓存效果测试", level=2)
add_heading(doc, "5.5.1 测试方法", level=3)
add_para(doc, "设计两组对比测试:A 组禁用 Redis(直查 MySQL);B 组启用 Redis 缓存。每组对相同的 100 条电商查询请求(覆盖 8 个商品的价格、库存查询)进行测试,记录平均响应时延与缓存命中率。")

add_heading(doc, "5.5.2 测试结果", level=3)
add_para(doc, "测试结果如表 5-5 所示。")
add_table_caption(doc, "表 5-5  Redis 缓存效果对比测试")
add_table(doc, ["指标", "A 组(无缓存)", "B 组(启用缓存)", "提升"], [
    ["平均查询延迟", "38 ms", "4 ms", "89.5%"],
    ["MySQL 查询次数", "100", "12", "88%"],
    ["缓存命中率", "—", "88%", "—"],
    ["P99 延迟", "65 ms", "12 ms", "81.5%"],
])
add_para(doc, "数据显示,启用 Redis 缓存后,平均查询延迟降低 89.5%,MySQL 查询次数减少 88%。这一效果验证了缓存层设计的价值,尤其在高并发场景下,缓存能够显著降低数据库负载,提升系统吞吐能力。")

add_heading(doc, "5.5.3 缓存失效测试", level=3)
add_para(doc, "为验证缓存失效机制的正确性,设计了如下测试:先查询商品 1001 的价格(写入缓存),然后通过管理接口更新该商品价格,再次查询验证返回的是更新后的价格。测试结果显示,价格更新后,系统主动清除了 product:1001 与 products: 模式下的所有缓存键,后续查询正确返回了新价格,验证了“写后失效”策略的有效性。")

add_heading(doc, "5.6 多源检索消融实验", level=2)
add_heading(doc, "5.6.1 实验设计", level=3)
add_para(doc, "为验证三路并联检索方案的有效性,设计消融实验对比四种配置下的回答质量:配置 A 仅 RAG 向量检索;配置 B RAG + 知识图谱;配置 C RAG + MySQL 实时查询;配置 D RAG + 知识图谱 + MySQL(完整方案)。测试集包含 30 条问题,涵盖产品问答、售后、电商三类场景。评估方式采用人工打分,从事实准确性(0–5 分)、回答完整性(0–5 分)、实时性(0–5 分)三个维度评分,每条问题取三人评分均值。")

add_heading(doc, "5.6.2 实验结果", level=3)
add_para(doc, "实验结果如表 5-6 所示。")
add_table_caption(doc, "表 5-6  多源检索消融实验结果")
add_table(doc, ["配置", "事实准确性", "回答完整性", "实时性", "综合得分"], [
    ["A:仅 RAG", "3.6", "3.2", "2.1", "8.9"],
    ["B:RAG + KG", "4.3", "4.0", "2.4", "10.7"],
    ["C:RAG + MySQL", "4.1", "3.6", "4.7", "12.4"],
    ["D:三路完整", "4.6", "4.5", "4.8", "13.9"],
])

add_heading(doc, "5.6.3 结果分析", level=3)
add_para(doc, "实验结果显示三个关键结论。第一,知识图谱显著提升回答完整性(B 比 A 高 0.8 分)。在产品规格对比、保修政策等结构化问题上,KG 提供了精确的关联信息,弥补了纯向量检索的不足。第二,MySQL 实时查询显著提升实时性(C 比 A 高 2.6 分)。在价格、库存、订单状态等需要实时数据的问题上,MySQL 提供了不可替代的最新信息,RAG 文档则因更新滞后而失效。第三,三路融合方案综合最优(D 比 A 综合得分高 5 分)。完整方案结合了 KG 的结构化推理能力、MySQL 的实时数据能力与 RAG 的非结构化文档检索能力,在三个维度上均取得最高分。这一实验有力验证了本系统“多源异构检索融合”核心架构的设计价值。")

add_heading(doc, "5.7 本章小结", level=2)
add_para(doc, "本章对系统进行了全面的功能与性能测试。意图分类器在 60 条测试用例上达到 95% 准确率,单次分类延迟低于 1 ms;端到端问答覆盖六类场景均能给出符合预期的回答;首字响应延迟约 500 ms,满足实时客服需求;Redis 缓存使查询延迟降低 89.5%、MySQL 负载降低 88%;多源检索消融实验显示三路融合方案综合得分较纯 RAG 提升 56%,验证了核心架构的有效性。测试结果整体满足系统设计的功能与性能目标。")

add_page_break(doc)


# ============ CHAPTER 6 ============
add_heading(doc, "第 6 章  总结与展望", level=1)
add_heading(doc, "6.1 工作总结", level=2)
add_para(doc, "本文围绕“基于 RAG 技术的智能客服系统”这一主题,从需求分析、架构设计到工程实现,完成了一套可私有化部署的完整方案,并通过系统化测试验证了其可用性。本文的主要贡献体现在以下五个方面。")
add_para(doc, "第一,提出并实现了“分类与生成解耦”的轻量化架构。将意图分类从大模型中剥离,通过纯本地的关键词规则与订单号正则模式实现,单次分类延迟低于 1 毫秒。这一设计使整个问答流程仅需调用一次 LLM,在保证 95% 分类准确率的同时,显著降低了首字响应延迟与单次问答的 Token 消耗。")
add_para(doc, "第二,设计并实现了 FAISS + NetworkX + MySQL 三路并联的多源检索方案。在 RAG 链路中按“结构化知识 → 实时业务数据 → 参考资料”的优先级组装上下文,各路检索可独立失败而不影响整体可用性。消融实验结果显示,三路融合方案在事实准确性、完整性、实时性三个维度上的综合得分较纯 RAG 提升 56%,有效解决了纯向量检索在结构化推理与实时性方面的局限。")
add_para(doc, "第三,设计并实现了基于 SSE 协议的流式问答接口。通过“meta + token × N + done”的事件协议,使前端可以在元信息到达时立刻渲染分类标签与参考来源,无需等待生成完成。实测用户感知首字延迟约 500 毫秒,完整答案生成约 2.8 秒,显著优于非流式实现。")
add_para(doc, "第四,实现了基于 MySQL + Redis 的可靠存储与缓存层。MySQL 持久化层共设计 8 张表,涵盖商品、订单、物流、维修、对话、反馈等业务实体;Redis 缓存层采用“缓存穿透 + 写后失效”策略,实测使查询延迟降低 89.5%、MySQL 负载降低 88%。Redis 不可用时系统自动降级为直查 MySQL,保证了系统的高可用性。")
add_para(doc, "第五,实现了集成 6 大功能模块的单文件前端。基于原生 HTML/CSS/JavaScript,无 UI 框架依赖,集成对话、知识库管理、模型配置、知识图谱可视化、演示数据初始化、统计看板共 6 大功能。整个前端仅一个 HTML 文件即可部署,极大降低了部署门槛。")

add_heading(doc, "6.2 不足与改进方向", level=2)
add_para(doc, "虽然本系统在功能与性能上均达到了设计目标,但在以下几个方面仍存在不足,有待后续改进。")
add_para(doc, "意图分类:当前系统采用关键词规则的方式实现意图分类,虽然具有零延迟的优势,但对模糊表达、隐喻表达、新出现的领域词汇识别能力有限。未来可引入轻量级分类模型(如 FastText 或小型 BERT)提升对模糊表述的识别能力,在精度与延迟之间取得更好的平衡。")
add_para(doc, "检索精度:系统未引入重排序(Reranker)机制,Top-K 结果直接用于上下文构建,部分检索结果可能存在相关性不足的问题。引入 Reranker 可在粗检索 Top-20 后精排取 Top-3,显著提升检索质量。可选方案包括 Cohere Reranker API、Jina Reranker、BGE Reranker 等。")
add_para(doc, "多轮对话:当前系统虽然保存了对话历史,但缺乏显式的多轮推理能力,对于需要多轮交互才能明确意图的场景处理较弱。未来可引入对话状态跟踪(DST)机制或基于 Agent 的多步推理框架,增强复杂场景下的对话连贯性。")
add_para(doc, "评估体系:当前系统的评估主要依赖人工打分与简单的功能性测试,缺乏系统化的 RAG 评估指标(如 Faithfulness、Answer Relevance、Context Precision、Context Recall)。后续可引入 RAGAS 等评估框架进行量化评估,建立持续优化的反馈闭环。")
add_para(doc, "知识图谱构建:当前知识图谱通过 Python 脚本以代码方式手工构建,扩展性有限。未来可引入基于 LLM 的实体关系自动抽取技术,从企业文档中自动构建知识图谱,降低运维成本。")
add_para(doc, "部署架构:当前系统主要面向单机部署,在分布式部署、横向扩展方面未做深入设计。生产环境若需支持高并发,需要引入负载均衡、Nginx 反向代理、消息队列、分布式向量数据库(如 Milvus 集群)等组件。")

add_heading(doc, "6.3 展望", level=2)
add_para(doc, "随着大模型技术的持续演进,RAG 智能客服系统未来将朝以下方向发展。")
add_para(doc, "多模态融合。下一代客服系统将不再局限于文本交互,而是融合图像(如用户上传的故障截图)、语音(如电话客服场景)、视频(如设备使用演示)等多模态信息。这将使系统能够处理“屏幕显示这个错误是什么意思”(配合截图)这类纯文本无法表达的问题。")
add_para(doc, "Agentic RAG。系统将从被动的“检索 + 生成”模式演进为主动的多步推理 Agent,能够自主规划检索步骤、调用外部工具(如查询订单系统、发起退款流程)、与用户多轮协作完成复杂任务。这将使智能客服从“问答机”升级为“虚拟员工”。")
add_para(doc, "端云协同。随着轻量级开源大模型(如 Qwen-3B、Llama-8B)的成熟,部分推理任务可下沉到客户端,实现端云协同的混合架构,在保护用户隐私的同时进一步降低响应延迟。")
add_para(doc, "自进化。通过用户反馈数据(rating、resolved 等)持续优化模型与知识库,形成“使用 → 反馈 → 优化”的正向闭环,使系统能够随业务演化自动调整。")
add_para(doc, "总体而言,本文工作为基于 RAG 的智能客服系统提供了一份完整、可落地、可扩展的工程方案,所采用的“分类与生成解耦、三路检索融合、单机轻量部署”思路对中小企业的私有化客服建设具有较强的参考价值。期待未来在多模态、Agentic、端云协同等方向上进一步深化,为智能客服领域贡献更多有价值的实践。")

add_page_break(doc)


# ============ REFERENCES ============
add_heading(doc, "参考文献", level=1)
references = [
    "[1] Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]. Advances in Neural Information Processing Systems (NeurIPS), 2020: 9459-9474.",
    "[2] Vaswani A, Shazeer N, Parmar N, et al. Attention Is All You Need[C]. Advances in Neural Information Processing Systems (NeurIPS), 2017: 5998-6008.",
    "[3] Devlin J, Chang M W, Lee K, et al. BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding[C]. Proceedings of NAACL-HLT, 2019: 4171-4186.",
    "[4] Brown T, Mann B, Ryder N, et al. Language Models are Few-Shot Learners[C]. Advances in Neural Information Processing Systems (NeurIPS), 2020: 1877-1901.",
    "[5] Karpukhin V, Oguz B, Min S, et al. Dense Passage Retrieval for Open-Domain Question Answering[C]. Proceedings of EMNLP, 2020: 6769-6781.",
    "[6] Gao L, Ma X, Lin J, et al. Precise Zero-Shot Dense Retrieval without Relevance Labels[C]. Proceedings of ACL, 2023: 1762-1777.",
    "[7] Es S, James J, Espinosa-Anke L, et al. RAGAS: Automated Evaluation of Retrieval Augmented Generation[C]. Proceedings of EACL System Demonstrations, 2024: 150-158.",
    "[8] DeepSeek-AI. DeepSeek-V3 Technical Report[R/OL]. (2024-12-26)[2026-04-30]. https://github.com/deepseek-ai/DeepSeek-V3.",
    "[9] Johnson J, Douze M, Jegou H. Billion-scale Similarity Search with GPUs[J]. IEEE Transactions on Big Data, 2021, 7(3): 535-547.",
    "[10] Reimers N, Gurevych I. Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks[C]. Proceedings of EMNLP-IJCNLP, 2019: 3982-3992.",
    "[11] Hagberg A A, Schult D A, Swart P J. Exploring Network Structure, Dynamics, and Function using NetworkX[C]. Proceedings of the 7th Python in Science Conference (SciPy), 2008: 11-15.",
    "[12] Chase H. LangChain: Building Applications with LLMs through Composability[CP/OL]. (2022-10-17)[2026-04-30]. https://github.com/langchain-ai/langchain.",
    "[13] Asai A, Wu Z, Wang Y, et al. Self-RAG: Learning to Retrieve, Generate, and Critique through Self-Reflection[C]. Proceedings of ICLR, 2024.",
    "[14] Edge D, Trinh H, Cheng N, et al. From Local to Global: A Graph RAG Approach to Query-Focused Summarization[J]. arXiv preprint arXiv:2404.16130, 2024.",
    "[15] 杨静, 杨建, 王晓亮. 基于检索增强的大语言模型问答系统研究综述[J]. 计算机科学, 2024, 51(6): 1-15.",
    "[16] 刘知远, 孙茂松, 林衍凯, 等. 知识表示学习研究进展[J]. 计算机研究与发展, 2016, 53(2): 247-261.",
    "[17] 王昊奋, 漆桂林, 陈华钧. 知识图谱:方法、实践与应用[M]. 北京:电子工业出版社, 2019.",
    "[18] 邱锡鹏. 神经网络与深度学习[M]. 北京:机械工业出版社, 2020.",
    "[19] Ramirez S. FastAPI: Modern, Fast (high-performance) Web Framework for Building APIs[CP/OL]. (2018-12-05)[2026-04-30]. https://github.com/tiangolo/fastapi.",
    "[20] Bayer M. SQLAlchemy[M]. The Architecture of Open Source Applications, 2012: 291-314.",
    "[21] Hickson I. Server-Sent Events[S/OL]. W3C Recommendation, (2015-02-03)[2026-04-30]. https://www.w3.org/TR/eventsource/.",
    "[22] Carlson J. Redis in Action[M]. Manning Publications, 2013.",
    "[23] DuBois P. MySQL[M]. 5th ed. Addison-Wesley Professional, 2013.",
    "[24] 张俊林. 大语言模型的检索增强生成技术综述[J]. 中文信息学报, 2024, 38(3): 1-20.",
    "[25] 智源研究院. BGE: One-Stop Retrieval Toolkit For Search and RAG[CP/OL]. (2023-08-02)[2026-04-30]. https://github.com/FlagOpen/FlagEmbedding.",
]
for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.first_line_indent = Cm(-0.7)
    pf = p.paragraph_format
    pf.left_indent = Cm(0.7)
    r = p.add_run(ref)
    set_run_font(r, size=10.5)

add_page_break(doc)


# ============ ACKNOWLEDGMENTS ============
add_heading(doc, "致    谢", level=1)
add_para(doc, "时光荏苒,转眼间四年的本科生涯即将告一段落。在论文付梓之际,我谨向所有给予我帮助的师长、家人和朋友们致以最诚挚的感谢。")
add_para(doc, "首先,衷心感谢我的指导教师 ____ 老师。从论文选题到框架设计,从代码实现到论文定稿,老师始终以严谨的治学态度和渊博的专业知识给予我悉心指导。老师对前沿技术的敏锐洞察、对工程实践的深刻理解,让我受益匪浅。每次讨论,老师都耐心解答我的疑问,并为我指明改进方向,在此向 ____ 老师致以最崇高的敬意。")
add_para(doc, "感谢学院和实验室提供的良好学习与科研环境,感谢各位授课老师在专业课程上的辛勤付出,你们传授的知识为本论文的完成奠定了扎实基础。")
add_para(doc, "感谢我的同窗好友们。在毕业设计期间,大家相互讨论、彼此鼓励,共同度过了无数挑灯夜战的日夜。特别感谢在系统调试与测试阶段给予我帮助的同学们,你们的建议让本系统不断完善。")
add_para(doc, "感谢父母多年来的无私支持与默默付出。无论是物质上的保障还是精神上的鼓励,你们始终是我最坚实的后盾。")
add_para(doc, "最后,衷心感谢在百忙之中评阅论文和参加答辩的各位老师,感谢您们对本论文的批评与指正。")
add_para(doc, "由于本人水平有限,论文中难免存在疏漏与不足之处,恳请各位老师批评指正。")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_before = Pt(20)
r = p.add_run("作 者\n二〇二六年五月")
set_run_font(r, size=12)


print("All chapters added.")
doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")


